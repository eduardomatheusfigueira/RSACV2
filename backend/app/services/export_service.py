#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Revsist — Serviço de Exportação de Dados e Relatórios Acadêmicos.
Gera planilhas Excel (.xlsx) com múltiplas abas, arquivos BibTeX (.bib)
e dados estruturados para o diagrama de fluxo PRISMA 2020.
"""

import io
import re
from typing import Dict

import pandas as pd
from sqlalchemy import or_
from sqlalchemy.orm import Session

from app.domain.afiliacao import e_nome_de_coletor
from app.infrastructure.persistence.models import (
    ExtractionAnswerModel,
    HarvestRunModel,
    PaperCriterionModel,
    PaperModel,
    ProjectModel,
    ProtocolModel,
)

# Caracteres com que o Excel e o LibreOffice reconhecem uma célula como
# fórmula. TAB e CR entram porque alguns leitores os tratam como separador e
# reposicionam o conteúdo seguinte no início de uma célula.
PREFIXOS_DE_FORMULA = ("=", "+", "-", "@", "\t", "\r")


def neutralizar_formula(valor):
    """
    Impede que um texto vindo de base externa seja executado como fórmula
    (doc 28 V-15, doc 29 §29.9.1).

    Títulos e resumos vêm de bases indexadas que o Revsist não controla. Um
    registro preparado com `=HYPERLINK(...)` ou `=cmd|...` vira execução na
    máquina de quem abrir a planilha exportada — e o pesquisador abre a
    planilha confiando que ela é produto do próprio trabalho.

    O apóstrofo à frente é a convenção que o Excel entende como "isto é texto":
    ele não aparece na célula, e o conteúdo é preservado na íntegra.
    """
    if not isinstance(valor, str) or not valor:
        return valor
    if valor.startswith(PREFIXOS_DE_FORMULA):
        return "'" + valor
    return valor


def sanitizar_nome_de_arquivo(nome: str, padrao: str = "exportacao") -> str:
    """
    Reduz um nome a caracteres seguros para `Content-Disposition`.

    O título do projeto era interpolado direto no cabeçalho: aspas ou quebra de
    linha ali quebram o cabeçalho, e o nome legível vai no `filename*` em
    UTF-8, que é onde acentos pertencem (§29.9.1).
    """
    import re

    limpo = re.sub(r"[^A-Za-z0-9._-]+", "_", nome or "").strip("._-")
    return limpo[:60] or padrao


def cabecalho_de_download(nome_legivel: str, extensao: str) -> dict:
    """Monta o `Content-Disposition` com nome sanitizado e nome legível."""
    from urllib.parse import quote

    ascii_seguro = f"{sanitizar_nome_de_arquivo(nome_legivel)}.{extensao}"
    legivel = quote(f"{nome_legivel}.{extensao}".replace("\n", " ").replace("\r", " "))
    return {
        "Content-Disposition": (
            f'attachment; filename="{ascii_seguro}"; filename*=UTF-8\'\'{legivel}'
        )
    }


class ExportService:
    """Serviço de Exportação Multi-formato."""

    @staticmethod
    def generate_excel(db: Session, project_id: str) -> io.BytesIO:
        """
        Gera um arquivo Excel (.xlsx) com 4 abas estruturadas:
          1. Artigos Incluídos
          2. Extração de Dados
          3. Artigos Excluídos
          4. Métricas PRISMA
        """
        project = db.query(ProjectModel).filter(ProjectModel.id == project_id).first()
        protocol = db.query(ProtocolModel).filter(ProtocolModel.project_id == project_id).first()

        # Duplicatas fora, pelo mesmo critério da fila de triagem e do
        # instantâneo: exportá-las inflaria toda contagem da planilha.
        all_papers = (
            db.query(PaperModel)
            .filter(
                PaperModel.project_id == project_id,
                or_(PaperModel.is_duplicate == False, PaperModel.is_duplicate.is_(None)),  # noqa: E712
            )
            .all()
        )
        included_papers = [p for p in all_papers if p.decision == "Incluído"]
        excluded_papers = [p for p in all_papers if p.decision == "Excluído"]
        pending_papers = [p for p in all_papers if p.decision == "Pendente"]

        # Mapa de critérios para lookup rápido
        criteria_by_id = {c.id: c for c in protocol.criteria} if protocol and protocol.criteria else {}

        # ── Aba 1: Artigos Incluídos ───────────────────────────────────
        inc_data = []
        for p in included_papers:
            sources_str = ", ".join(s.source_name for s in p.sources) if p.sources else ""
            evals = db.query(PaperCriterionModel).filter(PaperCriterionModel.paper_id == p.id, PaperCriterionModel.value == True).all()
            inc_crits = [criteria_by_id[e.criterion_id].text for e in evals if e.criterion_id in criteria_by_id and not criteria_by_id[e.criterion_id].is_exclusion]

            inc_data.append({
                "ID": p.id,
                "Título": p.title,
                "Autores": p.authors,
                "Orientador(a)": p.advisor or "",
                "Ano": p.year,
                "DOI": p.doi or "",
                "Periódico / Revista": p.journal or "",
                "Bases de Dados": sources_str,
                "Tipo de Pesquisa": p.research_type,
                # Vazio quando o campo traz só o nome do coletor: a procedência
                # já está em "Bases de Dados", e repeti-la sob o rótulo de
                # instituição faz a planilha afirmar uma afiliação que ninguém
                # coletou (doc 47 §B-01).
                "Instituição": "" if e_nome_de_coletor(p.institution) else p.institution,
                "URL / Link": p.download_url,
                "Critérios de Inclusão Atendidos": "; ".join(inc_crits),
                "Resumo (Abstract)": p.abstract,
                "Observações": p.observations,
            })
        df_included = pd.DataFrame(inc_data)

        # ── Aba 2: Extração de Dados ───────────────────────────────────
        ext_data = []
        questions = sorted(protocol.extraction_questions, key=lambda x: x.order) if protocol else []

        for p in included_papers:
            row = {
                "Paper ID": p.id,
                "Título": p.title,
                "Ano": p.year,
            }
            # Buscar respostas
            answers = {
                ans.question_id: ans.answer
                for ans in db.query(ExtractionAnswerModel).filter(ExtractionAnswerModel.paper_id == p.id).all()
            }
            for i, q in enumerate(questions):
                col_name = f"Q{i+1}: {q.text[:40]}..."
                row[col_name] = answers.get(q.id, "Não preenchido")
            ext_data.append(row)
        df_extraction = pd.DataFrame(ext_data)

        # ── Aba 3: Artigos Excluídos ───────────────────────────────────
        exc_data = []
        for p in excluded_papers:
            sources_str = ", ".join(s.source_name for s in p.sources) if p.sources else ""
            evals = db.query(PaperCriterionModel).filter(PaperCriterionModel.paper_id == p.id, PaperCriterionModel.value == True).all()
            exc_crits = [criteria_by_id[e.criterion_id].text for e in evals if e.criterion_id in criteria_by_id and criteria_by_id[e.criterion_id].is_exclusion]

            exc_data.append({
                "ID": p.id,
                "Título": p.title,
                "Autores": p.authors,
                "Ano": p.year,
                "DOI": p.doi or "",
                "Bases de Dados": sources_str,
                "Critérios de Exclusão Identificados": "; ".join(exc_crits),
                "Motivo / Observações": p.observations,
            })
        df_excluded = pd.DataFrame(exc_data)


        # ── Aba 4: Métricas PRISMA ─────────────────────────────────────
        runs = db.query(HarvestRunModel).filter(HarvestRunModel.project_id == project_id).all()
        total_found = sum(r.records_found for r in runs)
        total_dup = sum(r.records_duplicate for r in runs)

        prisma_data = [
            {"Etapa PRISMA 2020": "Total de Registros Identificados nas Bases", "Quantidade": total_found},
            {"Etapa PRISMA 2020": "Registros Duplicados Removidos", "Quantidade": total_dup},
            {"Etapa PRISMA 2020": "Registros Únicos a Triar (após deduplicação)", "Quantidade": len(all_papers)},
            # Triados são os que TÊM decisão. Contar o acervo inteiro aqui
            # declarava triagem que não aconteceu — ver `get_prisma_flow_data`.
            {
                "Etapa PRISMA 2020": "Registros Triados (Título e Resumo)",
                "Quantidade": len(included_papers) + len(excluded_papers),
            },
            {"Etapa PRISMA 2020": "Estudos Excluídos na Triagem 1", "Quantidade": len(excluded_papers)},
            {"Etapa PRISMA 2020": "Estudos Ainda Não Triados", "Quantidade": len(pending_papers)},
            {"Etapa PRISMA 2020": "Estudos Elegíveis / Incluídos na Síntese", "Quantidade": len(included_papers)},
        ]
        df_prisma = pd.DataFrame(prisma_data)

        # ── Aba 5: Indicadores Bibliométricos (doc 48 §7, doc 49 Fase 3)
        try:
            from app.services.bibliometria.indicadores import obter_indicadores_bibliometricos
            ind = obter_indicadores_bibliometricos(db, project_id, decision="Incluído")
            cagr = ind.get("production_temporal", {})
            bradford = ind.get("bradford", {})
            lotka = ind.get("lotka", {})
            colab = ind.get("collaboration", {})
            conc = ind.get("concentration", {})
            cit = ind.get("citations", {})
            oa = ind.get("open_access", {})

            biblio_data = [
                {"Dimensão": "Produção Temporal", "Indicador": "Período Analisado", "Valor": f"{cagr.get('year_start')} - {cagr.get('year_end')}" if cagr.get('year_start') else "—"},
                {"Dimensão": "Produção Temporal", "Indicador": "Taxa Composta de Crescimento Anual (CAGR)", "Valor": f"{cagr.get('cagr_pct')}%" if cagr.get('cagr_pct') is not None else "—"},
                {"Dimensão": "Produção Temporal", "Indicador": "Total de Estudos no Período", "Valor": cagr.get("total_period", 0)},
                {"Dimensão": "Lei de Bradford", "Indicador": "Total de Periódicos Identificados", "Valor": bradford.get("total_journals", 0)},
                {"Dimensão": "Lei de Bradford", "Indicador": "Razão de Periódicos por Zona (r1 : r2 : r3)", "Valor": bradford.get("formula_ratio", "—")},
                {"Dimensão": "Lei de Bradford", "Indicador": "Multiplicador de Bradford (k médio)", "Valor": bradford.get("k_multiplier") if bradford.get("k_multiplier") is not None else "—"},
                {"Dimensão": "Lei de Lotka", "Indicador": "Total de Autores Únicos", "Valor": lotka.get("n_authors", 0)},
                {"Dimensão": "Lei de Lotka", "Indicador": "Expoente da Lei de Potência (alpha)", "Valor": lotka.get("alpha") if lotka.get("alpha") is not None else "—"},
                {"Dimensão": "Lei de Lotka", "Indicador": "Estatística de Teste Kolmogorov-Smirnov (D_KS)", "Valor": lotka.get("d_ks") if lotka.get("d_ks") is not None else "—"},
                {"Dimensão": "Lei de Lotka", "Indicador": "Valor Crítico KS (5%)", "Valor": lotka.get("d_critical") if lotka.get("d_critical") is not None else "—"},
                {"Dimensão": "Lei de Lotka", "Indicador": "Veredicto de Aderência Formal", "Valor": lotka.get("p_verdict", "—")},
                {"Dimensão": "Colaboração", "Indicador": "Índice de Subramanyam (C)", "Valor": colab.get("subramanyam_index") if colab.get("subramanyam_index") is not None else "—"},
                {"Dimensão": "Colaboração", "Indicador": "Média de Autores por Artigo", "Valor": colab.get("avg_authors_per_paper", 0.0)},
                {"Dimensão": "Colaboração", "Indicador": "Artigos em Coautoria", "Valor": colab.get("multi_author_articles", 0)},
                {"Dimensão": "Colaboração", "Indicador": "Artigos com Autor Único", "Valor": colab.get("single_author_articles", 0)},
                {"Dimensão": "Concentração", "Indicador": "Coeficiente de Gini de Autores", "Valor": conc.get("gini_authors") if conc.get("gini_authors") is not None else "—"},
                {"Dimensão": "Concentração", "Indicador": "Coeficiente de Gini de Periódicos", "Valor": conc.get("gini_journals") if conc.get("gini_journals") is not None else "—"},
                {"Dimensão": "Concentração", "Indicador": "Índice Herfindahl-Hirschman (HHI) de Periódicos", "Valor": conc.get("hhi_journals") if conc.get("hhi_journals") is not None else "—"},
                {"Dimensão": "Impacto e Citações", "Indicador": "Índice h do Corpus", "Valor": cit.get("h_index", 0)},
                {"Dimensão": "Impacto e Citações", "Indicador": "Total de Citações Recebidas", "Valor": cit.get("total_citations", 0)},
                {"Dimensão": "Impacto e Citações", "Indicador": "Média de Citações por Artigo", "Valor": cit.get("mean_citations", 0.0)},
                {"Dimensão": "Impacto e Citações", "Indicador": "Mediana de Citações", "Valor": cit.get("median_citations", 0.0)},
                {"Dimensão": "Acesso Aberto", "Indicador": "Proporção Open Access", "Valor": f"{oa.get('open_access_pct')}%" if oa.get("open_access_pct") is not None else "—"},
                {"Dimensão": "Acesso Aberto", "Indicador": "Estudos em Acesso Aberto", "Valor": oa.get("open_access_count", 0)},
            ]
        except Exception:
            biblio_data = [{"Dimensão": "Erro", "Indicador": "Falha no cálculo", "Valor": "Indisponível"}]

        df_biblio = pd.DataFrame(biblio_data)

        # ── Aba 6: Metadados de Proveniência ───────────────────────────
        from datetime import datetime, timezone
        prov_data = [
            {"Propriedade": "Sistema", "Valor": "Revsist — Ambiente de Revisão Sistemática e Bibliometria"},
            {"Propriedade": "Versão do Motor", "Valor": "2.0.0"},
            {"Propriedade": "Projeto ID", "Valor": project_id},
            {"Propriedade": "Título do Projeto", "Valor": project.title},
            {"Propriedade": "Metodologia", "Valor": project.methodology},
            {"Propriedade": "Data da Exportação (UTC)", "Valor": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S")},
            {"Propriedade": "Total de Estudos no Projeto", "Valor": len(all_papers)},
            {"Propriedade": "Regra de Integridade Numérica", "Valor": "Cálculo 100% determinístico; nenhum indicador produzido por LLM (Doc 48 §2)"},
        ]
        df_prov = pd.DataFrame(prov_data)

        # Gerar BytesIO
        output = io.BytesIO()
        df_included = df_included.map(neutralizar_formula)
        df_extraction = df_extraction.map(neutralizar_formula)
        df_excluded = df_excluded.map(neutralizar_formula)
        df_prisma = df_prisma.map(neutralizar_formula)
        df_biblio = df_biblio.map(neutralizar_formula)
        df_prov = df_prov.map(neutralizar_formula)

        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            df_included.to_excel(writer, sheet_name="Artigos Incluídos", index=False)
            df_extraction.to_excel(writer, sheet_name="Extração de Dados", index=False)
            df_excluded.to_excel(writer, sheet_name="Artigos Excluídos", index=False)
            df_prisma.to_excel(writer, sheet_name="Métricas PRISMA 2020", index=False)
            df_biblio.to_excel(writer, sheet_name="Bibliometria", index=False)
            df_prov.to_excel(writer, sheet_name="Proveniência", index=False)

        output.seek(0)
        return output

    @staticmethod
    def generate_bibtex(db: Session, project_id: str, only_included: bool = True) -> str:
        """Gera arquivo BibTeX padronizado para gerenciadores de referências (Zotero, Mendeley, LaTeX)."""
        query = db.query(PaperModel).filter(PaperModel.project_id == project_id)
        if only_included:
            query = query.filter(PaperModel.decision == "Incluído")

        papers = query.all()
        bib_entries = []

        for p in papers:
            # Gerar chave de citação (ex: Smith2024Machine)
            first_author = p.authors.split(";")[0].split(",")[0].strip() if p.authors else "Unknown"
            first_author_clean = re.sub(r"[^a-zA-Z]", "", first_author)
            year_clean = p.year[:4] if p.year else "nodate"
            first_word = re.sub(r"[^a-zA-Z]", "", p.title.split()[0]) if p.title else "paper"
            cite_key = f"{first_author_clean}{year_clean}{first_word}"

            entry_type = "article" if "tese" not in p.research_type.lower() else "phdthesis"

            entry = f"@{entry_type}{{{cite_key},\n"
            entry += f"  title = {{{p.title}}},\n"
            if p.authors:
                entry += f"  author = {{{p.authors}}},\n"
            if p.year:
                entry += f"  year = {{{p.year}}},\n"
            if p.doi:
                entry += f"  doi = {{{p.doi}}},\n"
            if p.abstract:
                entry += f"  abstract = {{{p.abstract}}},\n"
            if p.download_url:
                entry += f"  url = {{{p.download_url}}},\n"
            entry += "}\n"
            bib_entries.append(entry)

        return "\n".join(bib_entries)

    @staticmethod
    def get_prisma_flow_data(db: Session, project_id: str) -> Dict:
        """Dados do diagrama de fluxo PRISMA 2020.

        **`records_screened` conta os registros que já foram triados** — os que
        têm decisão —, e não o tamanho do acervo.

        A distinção não é sutil e não é acadêmica. Antes, este campo devolvia
        `len(all_papers)`, o que declarava como "triados" todos os registros
        do projeto, inclusive os que ninguém tinha olhado. Medido nos acervos
        reais em 01/09/2026: um projeto com **454** estudos triados reportava
        16.578 (37×), e outro com **209** reportava 65.955 (316×).

        O diagrama PRISMA é o artefato regulado de uma revisão sistemática, e
        esse número entra nele, na planilha exportada e na tela. Errá-lo por
        duas ordens de grandeza é o pior defeito que este arquivo poderia ter.

        `records_to_screen` é o denominador — quantos entraram na fila — e
        existe para que a tela mostre "454 de 16.578" em vez de um número
        solto que se lê como se a triagem estivesse terminada.

        Duplicatas ficam de fora de todas as contagens, pelo mesmo critério da
        fila de triagem, do contador do projeto e do instantâneo. Antes, a
        consulta não as filtrava.
        """
        runs = db.query(HarvestRunModel).filter(HarvestRunModel.project_id == project_id).all()
        total_found = sum(r.records_found for r in runs)
        total_duplicates = sum(r.records_duplicate for r in runs)

        by_source = {}
        for r in runs:
            by_source[r.source_name] = by_source.get(r.source_name, 0) + r.records_found

        unicos = (
            db.query(PaperModel)
            .filter(
                PaperModel.project_id == project_id,
                or_(PaperModel.is_duplicate == False, PaperModel.is_duplicate.is_(None)),  # noqa: E712
            )
            .all()
        )
        included = [p for p in unicos if p.decision == "Incluído"]
        excluded = [p for p in unicos if p.decision == "Excluído"]
        pending = [p for p in unicos if p.decision == "Pendente"]

        return {
            "identification": {
                "total_records_identified": total_found,
                "sources_breakdown": by_source,
                "duplicates_removed": total_duplicates,
            },
            "screening": {
                # Quantos entraram na fila de triagem — o denominador.
                "records_to_screen": len(unicos),
                # Quantos de fato já foram triados: os que têm decisão.
                "records_screened": len(included) + len(excluded),
                "records_excluded": len(excluded),
                "records_pending": len(pending),
            },
            "included": {
                "studies_included_in_synthesis": len(included),
            },
        }

    @staticmethod
    def generate_search_log(db: Session, project_id: str, format_type: str = "json") -> tuple[bytes | str, str, str]:
        """
        Gera o Registro de Busca e Confronto Metodológico nos formatos DOCX, PDF, CSV e JSON (Doc 45 D-B).
        Retorna (conteudo, media_type, extensao).
        """
        import json
        import csv
        from app.infrastructure.persistence.models import (
            SearchStrategyModel,
            SearchExecutionModel,
            ProtocolVersionModel,
        )

        project = db.query(ProjectModel).filter(ProjectModel.id == project_id).first()
        protocol = db.query(ProtocolModel).filter(ProtocolModel.project_id == project_id).first()

        # 1. Recupera estratégias e execuções
        canonical_strat = None
        executions_list = []
        if protocol:
            canonical_strat = db.query(SearchStrategyModel).filter(
                SearchStrategyModel.protocol_id == protocol.id,
                SearchStrategyModel.kind == "canonica",
            ).first()

            execs = db.query(SearchExecutionModel).filter(
                SearchExecutionModel.protocol_id == protocol.id
            ).order_by(SearchExecutionModel.executed_at.asc()).all()

            for e in execs:
                executions_list.append({
                    "database": e.database,
                    "executed_at": e.executed_at.isoformat() if e.executed_at else "",
                    "query_sent": e.query_sent,
                    "filters": json.loads(e.filters) if e.filters else {},
                    "records_returned": e.records_returned,
                    "records_after_dedup": e.records_after_dedup,
                    "error": e.error or "",
                })

        # Fallback sintético a partir de HarvestRuns se não houver registros em SearchExecutionModel
        if not executions_list:
            runs = db.query(HarvestRunModel).filter(HarvestRunModel.project_id == project_id).all()
            for r in runs:
                executions_list.append({
                    "database": r.source_name,
                    "executed_at": r.created_at.isoformat() if r.created_at else "",
                    "query_sent": f"Execução de busca na base {r.source_name}",
                    "filters": json.loads(r.filters_applied) if r.filters_applied else {},
                    "records_returned": r.records_found,
                    "records_after_dedup": max(0, r.records_found - r.records_duplicate),
                    "error": "",
                })

        # Dados estruturados base
        log_data = {
            "project_id": project_id,
            "project_title": project.title if project else "Sem título",
            "review_design": protocol.review_design if protocol else "D4",
            "reporting_guideline": protocol.reporting_guideline if protocol else "PRISMA-ScR",
            "protocol_version": protocol.current_version if protocol else "v1.0",
            "canonical_query": canonical_strat.rendered_query if canonical_strat else "",
            "canonical_blocks": json.loads(canonical_strat.blocks) if canonical_strat and canonical_strat.blocks else [],
            "executions": executions_list,
        }

        # ── Formato JSON ────────────────────────────────────────────────
        if format_type.lower() == "json":
            content_str = json.dumps(log_data, indent=2, ensure_ascii=False)
            return content_str, "application/json", "json"

        # ── Formato CSV ─────────────────────────────────────────────────
        if format_type.lower() == "csv":
            output = io.StringIO()
            writer = csv.writer(output)
            writer.writerow([
                "Base de Dados",
                "Data/Hora de Execucao",
                "Consulta Enviada",
                "Filtros Aplicados",
                "Registros Brutos",
                "Apos Deduplicacao",
                "Erros / Observacoes",
            ])
            for e in executions_list:
                writer.writerow([
                    neutralizar_formula(e["database"]),
                    e["executed_at"],
                    neutralizar_formula(e["query_sent"]),
                    neutralizar_formula(json.dumps(e["filters"], ensure_ascii=False)),
                    e["records_returned"],
                    e["records_after_dedup"],
                    neutralizar_formula(e["error"]),
                ])
            return output.getvalue(), "text/csv; charset=utf-8", "csv"

        # ── Formato DOCX ────────────────────────────────────────────────
        if format_type.lower() == "docx":
            import docx
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT

            doc = docx.Document()

            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_after = Pt(6)
            title_run = title_p.add_run("Registro de Busca e Confronto Metodológico (PRISMA-S)")
            title_run.bold = True
            title_run.font.size = Pt(16)
            title_run.font.color.rgb = RGBColor(15, 23, 42)

            sub_p = doc.add_paragraph()
            sub_p.paragraph_format.space_after = Pt(14)
            sub_run = sub_p.add_run(f"Projeto: {log_data['project_title']} | Versão do Protocolo: {log_data['protocol_version']}")
            sub_run.font.size = Pt(10)
            sub_run.font.italic = True
            sub_run.font.color.rgb = RGBColor(71, 85, 105)

            # Seção 1: Metadados do Protocolo
            h1 = doc.add_heading("1. Metadados do Protocolo de Pesquisa", level=1)
            meta_p = doc.add_paragraph()
            meta_p.add_run(f"Desenho Metodológico: ").bold = True
            meta_p.add_run(f"{log_data['review_design']}\n")
            meta_p.add_run(f"Diretriz de Relato: ").bold = True
            meta_p.add_run(f"{log_data['reporting_guideline']}\n")
            meta_p.add_run(f"Estratégia Canônica: ").bold = True
            meta_p.add_run(f"{log_data['canonical_query'] or 'Não configurada'}\n")

            # Seção 2: Confronto de Execuções por Base
            doc.add_heading("2. Registro de Execução e Coleta nas Bases de Dados", level=1)
            
            table = doc.add_table(rows=1, cols=5)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells
            headers = ["Base", "Data/Hora", "Consulta Enviada", "Registros", "Pós-Dedup"]
            for idx, h_text in enumerate(headers):
                hdr_cells[idx].text = h_text
                for p in hdr_cells[idx].paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(9)

            for e in executions_list:
                row_cells = table.add_row().cells
                row_cells[0].text = e["database"]
                row_cells[1].text = e["executed_at"][:16].replace("T", " ") if e["executed_at"] else "-"
                row_cells[2].text = e["query_sent"][:120] + ("..." if len(e["query_sent"]) > 120 else "")
                row_cells[3].text = str(e["records_returned"])
                row_cells[4].text = str(e["records_after_dedup"])
                for c in row_cells:
                    for p in c.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(8.5)

            stream = io.BytesIO()
            doc.save(stream)
            stream.seek(0)
            return stream.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"

        # ── Formato PDF ─────────────────────────────────────────────────
        if format_type.lower() == "pdf":
            from reportlab.lib.pagesizes import A4
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

            stream = io.BytesIO()
            doc_pdf = SimpleDocTemplate(
                stream,
                pagesize=A4,
                rightMargin=36,
                leftMargin=36,
                topMargin=36,
                bottomMargin=36,
            )

            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                "TitleStyle",
                parent=styles["Heading1"],
                fontSize=14,
                leading=16,
                textColor=colors.HexColor("#0f172a"),
            )
            body_style = ParagraphStyle(
                "BodyStyle",
                parent=styles["Normal"],
                fontSize=8.5,
                leading=11,
                textColor=colors.HexColor("#334155"),
            )
            table_hdr_style = ParagraphStyle(
                "TableHdr",
                parent=styles["Normal"],
                fontSize=8,
                leading=9,
                textColor=colors.white,
                fontName="Helvetica-Bold",
            )
            table_cell_style = ParagraphStyle(
                "TableCell",
                parent=styles["Normal"],
                fontSize=7.5,
                leading=9,
                textColor=colors.HexColor("#1e293b"),
            )

            story = []
            story.append(Paragraph("Registro de Busca e Confronto Metodológico (PRISMA-S)", title_style))
            story.append(Spacer(1, 4))
            story.append(Paragraph(f"Projeto: {log_data['project_title']} | Versão do Protocolo: {log_data['protocol_version']}", body_style))
            story.append(Spacer(1, 10))

            # Resumo Metodológico
            meta_text = (
                f"<b>Desenho:</b> {log_data['review_design']} | "
                f"<b>Diretriz de Relato:</b> {log_data['reporting_guideline']}<br/>"
                f"<b>Consulta Canônica:</b> <i>{log_data['canonical_query'] or 'Definida por pares de termos'}</i>"
            )
            story.append(Paragraph(meta_text, body_style))
            story.append(Spacer(1, 10))

            # Tabela de Execuções
            story.append(Paragraph("<b>Execuções de Busca por Base de Dados</b>", styles["Heading3"]))
            story.append(Spacer(1, 6))

            table_data = [
                [
                    Paragraph("Base", table_hdr_style),
                    Paragraph("Data/Hora", table_hdr_style),
                    Paragraph("Consulta Enviada", table_hdr_style),
                    Paragraph("Brutos", table_hdr_style),
                    Paragraph("Pós-Dedup", table_hdr_style),
                ]
            ]

            for e in executions_list:
                dt_str = e["executed_at"][:16].replace("T", " ") if e["executed_at"] else "-"
                q_text = e["query_sent"][:100] + ("..." if len(e["query_sent"]) > 100 else "")
                table_data.append([
                    Paragraph(e["database"], table_cell_style),
                    Paragraph(dt_str, table_cell_style),
                    Paragraph(q_text, table_cell_style),
                    Paragraph(str(e["records_returned"]), table_cell_style),
                    Paragraph(str(e["records_after_dedup"]), table_cell_style),
                ])

            pdf_table = Table(table_data, colWidths=[65, 80, 240, 50, 50])
            pdf_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e293b")),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(pdf_table)

            doc_pdf.build(story)
            stream.seek(0)
            return stream.getvalue(), "application/pdf", "pdf"

        # Fallback default
        return json.dumps(log_data), "application/json", "json"

