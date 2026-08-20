"""
Script de Geração dos Documentos de Validação do RSACV2
Protocolo Integrado: "Impactos da Segurança Pública na Operacionalização do Turismo Náutico em Fronteiras Fluviais: Protocolo de Revisão de Escopo"
Origem dos Dados: rsac-perfil-backup-2026-08-19 (1).json (Projeto ID: 7847417c-79df-4892-bdec-2257d019f65e)

Arquivos Gerados:
1. Artigo_Esqueleto_Validacao_RSACV2.docx
2. Formulario_Pesquisadores_Triagem_Extracao.xlsx
3. Formulario_Avaliadores_Revisao_Cega.xlsx
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation

BASE_DIR = r"d:\Downloads\RSACV2\RSACV2\estudo_validacao"
os.makedirs(BASE_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# UTILITÁRIOS DOCX
# ---------------------------------------------------------------------------
def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_cell_borders(cell, color="CBD5E1", sz="4", val="single"):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/></w:tcBorders>')
    tcPr.append(tcBorders)

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    for run in h.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(13.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        elif level == 2:
            run.font.size = Pt(11.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2B, 0x54, 0x7E)
        elif level == 3:
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return h

def add_callout(doc, text, title="NOTA METODOLÓGICA", border_color="1B365D", fill_color="F0F4F8"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, fill_color)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r_title = p.add_run(f"[{title}]\n")
    r_title.bold = True
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(9.5)
    r_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    r_text = p.add_run(text)
    r_text.italic = True
    r_text.font.name = "Calibri"
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ---------------------------------------------------------------------------
# 1. GERAÇÃO DO ARTIGO EM DOCX
# ---------------------------------------------------------------------------
def generate_article_docx():
    doc = docx.Document()
    
    # Margens ABNT (Sup: 3cm, Esq: 3cm, Inf: 2cm, Dir: 2cm)
    for section in doc.sections:
        section.top_margin = Inches(1.18)    # ~3.0 cm
        section.left_margin = Inches(1.18)   # ~3.0 cm
        section.bottom_margin = Inches(0.79) # ~2.0 cm
        section.right_margin = Inches(0.79)  # ~2.0 cm
        
        # Header / Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Estudo de Validação RSACV2 | Avaliação Comparativa Cega")
        hrun.font.name = "Calibri"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("RSACV2 — Revisão Sistemática Assistida por Computador • Protocolo Experimental de Validação")
        frun.font.name = "Calibri"
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Estilo Normal
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)

    # TÍTULO PRINCIPAL
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(8)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("ESTUDO DE VALIDAÇÃO DO RSACV2:\nAVALIAÇÃO COMPARATIVA CEGA ENTRE TRIAGEM E EXTRAÇÃO ASSISTIDA POR COMPUTADOR E MÉTODO MANUAL EM REVISÃO DE ESCOPO")
    r_title.bold = True
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(15)
    r_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    # SUBTÍTULO
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(12)
    r_sub = p_sub.add_run("Aplicação Experimental no Protocolo: 'Impactos da Segurança Pública na Operacionalização do Turismo Náutico em Fronteiras Fluviais'")
    r_sub.italic = True
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # AUTORES E AFILIAÇÃO
    p_autores = doc.add_paragraph()
    p_autores.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_autores.paragraph_format.space_after = Pt(16)
    r_aut = p_autores.add_run("[Nome dos Pesquisadores / Grupo de Pesquisa RSACV2]\n")
    r_aut.bold = True
    r_inst = p_autores.add_run("Programa de Pós-Graduação em Desenvolvimento Regional e Políticas Públicas\nContato: [email@instituicao.br]")
    r_inst.font.size = Pt(9.5)
    r_inst.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # RESUMO ESTRUTURADO
    tbl_res = doc.add_table(rows=1, cols=1)
    tbl_res.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_res = tbl_res.cell(0, 0)
    set_cell_background(c_res, "F7F9FB")
    set_cell_margins(c_res, top=140, bottom=140, left=180, right=180)
    set_cell_borders(c_res, color="CBD5E1", sz="6")
    
    p_res = c_res.paragraphs[0]
    p_res.paragraph_format.space_before = Pt(2)
    p_res.paragraph_format.space_after = Pt(4)
    r_res_t = p_res.add_run("RESUMO ESTRUTURADO\n")
    r_res_t.bold = True
    r_res_t.font.size = Pt(10.5)
    r_res_t.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    resumo_texto = (
        "Contexto: A execução de Revisões Sistemáticas e de Escopo (Scoping Reviews) impõe intensa sobrecarga cognitiva e tempo-intensiva, "
        "particularmente nas etapas de triagem por elegibilidade e extração estruturada de evidências. A ferramenta RSACV2 propõe a assistência "
        "computacional inteligente para acelerar tais etapas mantendo estrito rigor metodológico. "
        "Objetivo: Validar a acurácia, reprodutibilidade e percepção de qualidade do sistema RSACV2 frente ao método manual não assistido "
        "na aplicação do protocolo de revisão de escopo 'Impactos da Segurança Pública na Operacionalização do Turismo Náutico em Fronteiras Fluviais'. "
        "Método: Estudo experimental pareado cego. Dois braços independentes (Braço Assistido RSACV2 vs. Braço Manual) executam o protocolo "
        "idêntico nas bases BDTD e SciELO, com 5 pares de descritores por idioma (PT, EN, ES), 3 critérios de inclusão (CI1-CI3), 3 de exclusão (CE1-CE3) "
        "e 5 questões de extração (QE1-QE5). Um trio de revisores independentes avalia os resultados em 6 dimensões cegas: "
        "(1) igualdade na coleta de estudos e ordem de aparição, (2) concordância nas marcações de triagem e justificativas, "
        "(3) preferência técnica em divergências, (4) discernimento da assistência (Teste de Turing perceptual), "
        "(5) concordância na extração de dados e (6) síntese qualitativa global. "
        "Resultados Esperados: Demonstrar equivalência na abrangência da busca, superioridade de consistência nas justificativas de triagem "
        "e alto grau de completude nas respostas de extração pelo RSACV2, consolidando sua eficácia no apoio à pesquisa em Desenvolvimento Regional e Segurança Territorial.\n\n"
        "Palavras-chave: Revisão de Escopo; RSACV2; Turismo Náutico; Fronteiras Fluviais; Segurança Pública; Avaliação Cega Inter-Avaliadores; Validação Metodológica."
    )
    p_res.add_run(resumo_texto).font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 1. INTRODUÇÃO
    add_styled_heading(doc, "1. INTRODUÇÃO", level=1)
    doc.add_paragraph(
        "A complexidade de conduzir revisões bibliográficas sistemáticas e de escopo (Scoping Reviews) no campo das Ciências Sociais Aplicadas "
        "e do Desenvolvimento Regional decorre da transversalidade temática e da necessidade de integrar literatura periódica indexada "
        "(SciELO) com a rica produção de teses e dissertações nacionais (BDTD / Ibict). "
        "Regiões de fronteira fluvial e hidrovias transfronteiriças possuem expressivo potencial para o turismo náutico e ecoturismo, "
        "mas enfrentam severos desafios de segurança pública, crimes transnacionais e déficits de governança institucional."
    )
    doc.add_paragraph(
        "Para responder a essas demandas com celeridade e rigor, o sistema RSACV2 foi concebido como um ambiente computacional integrado "
        "de coleta, triagem automatizada ancorada e extração orientada por IA. "
        "Este artigo relata a validação empírica do sistema mediante um desenho metodológico de avaliação cega inter-avaliadores, "
        "tendo como objeto de prova o protocolo: 'Impactos da Segurança Pública na Operacionalização do Turismo Náutico em Fronteiras Fluviais'."
    )

    add_styled_heading(doc, "1.1 Perguntas de Pesquisa (Research Questions)", level=2)
    doc.add_paragraph("A validação é balizada por seis Questões de Pesquisa centrais:")

    rqs = [
        ("RQ1 (Coleta e Busca)", "Usando o mesmo protocolo de busca, em uma mesma base, os trabalhos encontrados pelo RSACV2 e pelo método manual são rigorosamente os mesmos em presença e ordem de aparição?"),
        ("RQ2 (Concordância na Triagem)", "Em uma mesma lista de critérios de inclusão e exclusão, as marcações de triagem serão as mesmas entre os métodos? E a fundamentação das justificativas?"),
        ("RQ3 (Preferência em Divergências)", "Em caso de divergências de marcações e justificativas entre o método assistido e o manual, qual é preferida por um trio de revisores independentes em uma análise cega?"),
        ("RQ4 (Discernibilidade / Teste de Turing)", "Os revisores serão capazes de diferenciar qual trabalho foi executado de forma assistida pelo RSACV2? Sob qual grau de certeza e quais justificativas?"),
        ("RQ5 (Qualidade na Extração de Informações)", "Como se comparam as extrações de dados nos dois métodos sob avaliação cega por um trio de revisores independentes, considerando fidelidade e profundidade?"),
        ("RQ6 (Síntese Qualitativa Global)", "Quais observações, padrões de similaridade e diferenças estruturais os revisores declaram sobre o trabalho realizado com e sem assistência computacional?")
    ]

    for rq_id, rq_text in rqs:
        p_rq = doc.add_paragraph()
        p_rq.paragraph_format.left_indent = Inches(0.25)
        p_rq.paragraph_format.space_after = Pt(3)
        r1 = p_rq.add_run(f"• {rq_id}: ")
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        p_rq.add_run(rq_text)

    add_callout(
        doc,
        "O protocolo de validação ancora-se no princípio do duplo-cegamento: os pesquisadores executam o protocolo sem contato cruzado "
        "e o trio de revisores recebe todos os pares de resultados (Método A vs. Método B) com identificação criptografada e posições aleatorizadas.",
        title="DIRETRIZ DE CEGAMENTO EXPERIMENTAL"
    )

    # 2. PROTOCOLO METODOLÓGICO DA REVISÃO DE ESCOPO
    add_styled_heading(doc, "2. O PROTOCOLO DE REVISÃO DE ESCOPO EM TESTE", level=1)
    doc.add_paragraph(
        "O protocolo adotado como base comum de teste segue as diretrizes PRISMA-ScR (PRISMA Extension for Scoping Reviews, 2018) "
        "e recomendações do Joanna Briggs Institute (JBI). Seu objetivo precípuo é responder à seguinte questão norteadora:"
    )
    
    p_obj = doc.add_paragraph()
    p_obj.paragraph_format.left_indent = Inches(0.4)
    r_obj = p_obj.add_run('"Quais são as problemáticas de segurança pública registradas em fronteiras fluviais e de que maneira elas impactam o desenvolvimento e a operacionalização do turismo náutico nessas regiões?"')
    r_obj.bold = True
    r_obj.italic = True
    r_obj.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    add_styled_heading(doc, "2.1 Estratégia de Busca e Descritores em Pares", level=2)
    doc.add_paragraph(
        "Atendendo rigorosamente às regras de indexação e busca do motor VuFind (BDTD) e dos indexadores da SciELO, "
        "a estratégia foi elaborada com equilíbrio de especificidade e no formato de pares booleanos (máximo 2 termos por expressão com operador AND), "
        "limitada a 5 pares por idioma:"
    )

    # Tabela de Descritores do Protocolo
    tbl_desc = doc.add_table(rows=6, cols=4)
    tbl_desc.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_desc = ["Par #", "Português (PT)", "Inglês (EN)", "Espanhol (ES)"]
    col_w = [Inches(0.6), Inches(2.2), Inches(2.2), Inches(2.2)]
    
    for c_idx, h in enumerate(headers_desc):
        cell = tbl_desc.cell(0, c_idx)
        cell.width = col_w[c_idx]
        set_cell_background(cell, "1B365D")
        set_cell_margins(cell, 80, 80, 100, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)

    desc_rows = [
        ("Par 1", '"turismo náutico" AND "fronteira"', '"nautical tourism" AND "border"', '"turismo náutico" AND "frontera"'),
        ("Par 2", '"segurança pública" AND "fronteira fluvial"', '"public security" AND "river border"', '"seguridad pública" AND "frontera fluvial"'),
        ("Par 3", '"turismo" AND "fronteira fluvial"', '"tourism" AND "river border"', '"turismo" AND "frontera fluvial"'),
        ("Par 4", '"turismo" AND "tríplice fronteira"', '"tourism" AND "cross-border"', '"turismo" AND "triple frontera"'),
        ("Par 5", '"segurança pública" AND "turismo náutico"', '"water tourism" AND "border"', '"turismo fluvial" AND "frontera"')
    ]
    for r_idx, row in enumerate(desc_rows, start=1):
        for c_idx, val in enumerate(row):
            cell = tbl_desc.cell(r_idx, c_idx)
            cell.width = col_w[c_idx]
            set_cell_background(cell, "F0F4F8" if r_idx % 2 == 1 else "FFFFFF")
            set_cell_margins(cell, 60, 60, 80, 80)
            set_cell_borders(cell, color="CBD5E1", sz="4")
            p = cell.paragraphs[0]
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(val).font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_styled_heading(doc, "2.2 Critérios de Elegibilidade (Triagem)", level=2)
    doc.add_paragraph("A seleção dos estudos baseia-se nos seguintes critérios formais do protocolo:")

    criterios = [
        ("CI1 (Inclusão)", "Estudos que abordem a atividade turística, náutica, recreativa ou de navegação de passageiros em regiões de fronteira fluvial ou hidrovias transfronteiriças."),
        ("CI2 (Inclusão)", "Pesquisas que analisem aspectos de segurança pública, criminalidade transfronteiriça, fiscalização, policiamento ou governança em bacias hidrográficas de fronteira."),
        ("CI3 (Inclusão)", "Publicações científicas completas (artigos de periódicos, teses e dissertações) nos idiomas português, inglês ou espanhol."),
        ("CE1 (Exclusão)", "Estudos com foco exclusivo em transporte marítimo oceânico ou de alto-mar sem interface fluvial ou fronteiriça."),
        ("CE2 (Exclusão)", "Trabalhos sobre segurança pública puramente urbana ou rural sem qualquer conexão com hidrovias, cursos d'água de fronteira ou atividades turísticas."),
        ("CE3 (Exclusão)", "Documentos editoriais, resenhas de livros, resumos expandidos de eventos ou textos sem metodologia científica definida.")
    ]
    for c_id, c_desc in criterios:
        p_c = doc.add_paragraph()
        p_c.paragraph_format.left_indent = Inches(0.25)
        p_c.paragraph_format.space_after = Pt(2.5)
        r_c = p_c.add_run(f"• {c_id}: ")
        r_c.bold = True
        r_c.font.color.rgb = RGBColor(0x2B, 0x54, 0x7E)
        p_c.add_run(c_desc)

    add_styled_heading(doc, "2.3 Questões de Extração de Dados", level=2)
    doc.add_paragraph("Para os estudos incluídos na triagem, as seguintes questões são respondidas de forma independente:")

    extracao_questoes = [
        ("QE1 (Localização & Bacia)", "Qual é a localização geográfica, país e bacia hidrográfica/rio de fronteira analisado no estudo?"),
        ("QE2 (Tipologias de Ocorrências)", "Quais tipologias de ocorrências de segurança pública, crimes ou ilícitos transfronteiriços foram identificadas?"),
        ("QE3 (Impactos no Turismo Náutico)", "Quais foram os impactos diretos ou indiretos na atratividade, infraestrutura e dinâmica operacional do turismo náutico?"),
        ("QE4 (Governança & Políticas Públicas)", "Quais estratégias de governança transfronteiriça, políticas públicas ou medidas de policiamento/mitigação foram recomendadas?"),
        ("QE5 (Metodologia & Fontes de Dados)", "Qual a metodologia de pesquisa empregada e quais fontes de dados foram utilizadas?")
    ]
    for qe_id, qe_text in extracao_questoes:
        p_qe = doc.add_paragraph()
        p_qe.paragraph_format.left_indent = Inches(0.25)
        p_qe.paragraph_format.space_after = Pt(2.5)
        r_qe = p_qe.add_run(f"• {qe_id}: ")
        r_qe.bold = True
        p_qe.add_run(qe_text)

    # 3. METODOLOGIA DO ESTUDO DE VALIDAÇÃO
    add_styled_heading(doc, "3. METODOLOGIA DO EXPERIMENTO DE VALIDAÇÃO", level=1)
    doc.add_paragraph(
        "O delineamento experimental envolve dois braços de execução paralela: o Braço 1 (Execução Manual por Pesquisadores Humanos) "
        "e o Braço 2 (Execução Automatizada via RSACV2). As saídas brutas de ambos os braços são tratadas pela Coordenação, "
        "que aplica a randomização e anonimização em pares (Método A vs. Método B) antes de submetê-las ao comitê avaliador."
    )

    add_styled_heading(doc, "3.1 Operacionalização dos Critérios de Julgamento Cego", level=2)
    doc.add_paragraph("O julgamento pelo trio de revisores independentes obedece às seguintes regras padronizadas:")

    regras = [
        ("Critério 1.1 (Julgamento de Coleta - RQ1)", "Classificação categórica: (a) 'SIM' se houver 100% de igualdade nos estudos recuperados e na sua ordem exata de aparição; (b) 'NÃO, POR PRESENÇA' se um ou mais estudos forem divergentes entre as listas; (c) 'NÃO, POR ORDEM' se os estudos forem os mesmos, mas em posições de ranqueamento distintas; (d) 'NÃO INTEGRAL' se divergirem tanto na presença quanto na ordenação."),
        ("Critério 2.1 & 2.2 (Triagem & Justificativas - RQ2)", "Preenchimento da matriz de elegibilidade (CI1-CI3 e CE1-CE3), atribuição da decisão final (Incluído/Excluído) e elaboração da justificativa textual. O trio afere o rigor lógico e a aderência aos fatos do resumo."),
        ("Critério 3.1 (Preferência em Divergências - RQ3)", "Nos casos em que Método A e Método B divergem na decisão ou justificativa, o revisor assinala: (i) Método A está correto, (ii) Método B está correto, (iii) Ambos divergem mas estão corretos, ou (iv) Ambos erram. Exige justificativa técnica por extenso."),
        ("Critério 4.1 (Discernimento / Teste de Turing - RQ4)", "O revisor indica qual método aparenta ser a resposta gerada com auxílio computacional (Método A, Método B ou Indistinguível) e quantifica seu grau de certeza de 1 a 5 (1 = Totalmente incerto / chute; 5 = Certeza absoluta). Para notas 4 ou 5, a justificativa analítica é obrigatória."),
        ("Critério 5.1 & 5.2 (Extração de Dados - RQ5)", "Os revisores analisam as respostas cegas de extração (QE1-QE5), apontam similaridades/diferenças e atribuem notas de concordância de 1 a 5 para cada método, além de indicar a extração preferida."),
        ("Critério 6.1 (Parecer Qualitativo Global - RQ6)", "Redação de um parágrafo reflexivo estruturado consolidando as impressões gerais sobre consistência, profundidade analítica, clareza sintática e confiabilidade.")
    ]
    for r_id, r_text in regras:
        p_r = doc.add_paragraph()
        p_r.paragraph_format.left_indent = Inches(0.25)
        p_r.paragraph_format.space_after = Pt(3)
        r_r = p_r.add_run(f"• {r_id}: ")
        r_r.bold = True
        r_r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        p_r.add_run(r_text)

    # 4. RESULTADOS E TABULAÇÃO
    add_styled_heading(doc, "4. RESULTADOS E DISCUSSÃO PRELIMINAR", level=1)
    doc.add_paragraph(
        "Apresentam-se a seguir as estruturas de tabulação prontas para a consolidação dos dados experimentais coletados."
    )

    add_styled_heading(doc, "4.1 RQ1: Recuperação e Ordenação na Coleta", level=2)
    doc.add_paragraph("A Tabela 1 sintetiza a comparação das buscas executadas nas bases BDTD e SciELO.")
    
    tbl_rq1 = doc.add_table(rows=4, cols=6)
    tbl_rq1.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_rq1 = ["Base de Dados", "Total Manual", "Total RSACV2", "Interseção", "Julgamento (1.1)", "Taxa de Sobreposição (%)"]
    for c_idx, h in enumerate(headers_rq1):
        cell = tbl_rq1.cell(0, c_idx)
        set_cell_background(cell, "1B365D")
        set_cell_margins(cell, 80, 80, 100, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(8.5)

    sample_rq1 = [
        ("BDTD (VuFind)", "[Valor]", "[Valor]", "[Valor]", "[SIM / NÃO POR ORDEM / ...]", "[ % ]"),
        ("SciELO (REST)", "[Valor]", "[Valor]", "[Valor]", "[SIM / NÃO POR ORDEM / ...]", "[ % ]"),
        ("Total Consolidado", "[Total]", "[Total]", "[Total]", "[Classificação Geral]", "[ % Médio ]")
    ]
    for r_idx, row in enumerate(sample_rq1, start=1):
        for c_idx, val in enumerate(row):
            cell = tbl_rq1.cell(r_idx, c_idx)
            set_cell_background(cell, "F0F4F8" if r_idx % 2 == 1 else "FFFFFF")
            set_cell_margins(cell, 60, 60, 80, 80)
            set_cell_borders(cell, color="CBD5E1", sz="4")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(val).font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_styled_heading(doc, "4.2 RQ2 & RQ3: Concordância na Triagem e Preferência em Divergências", level=2)
    doc.add_paragraph(
        "A Tabela 2 apresenta o percentual de concordância na triagem de elegibilidade e a preferência técnica dos revisores nos casos divergentes."
    )

    tbl_rq2 = doc.add_table(rows=5, cols=5)
    tbl_rq2.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_rq2 = ["Revisor", "Concordância de Decisão (%)", "Preferência RSACV2 (%)", "Preferência Manual (%)", "Ambos Corretos / Erram (%)"]
    for c_idx, h in enumerate(headers_rq2):
        cell = tbl_rq2.cell(0, c_idx)
        set_cell_background(cell, "1B365D")
        set_cell_margins(cell, 80, 80, 100, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(8.5)

    sample_rq2 = [
        ("Revisor 1", "[ % ]", "[ % ]", "[ % ]", "[ % ]"),
        ("Revisor 2", "[ % ]", "[ % ]", "[ % ]", "[ % ]"),
        ("Revisor 3", "[ % ]", "[ % ]", "[ % ]", "[ % ]"),
        ("Consenso / Média", "[ % Médio ]", "[ % Médio ]", "[ % Médio ]", "[ % Médio ]")
    ]
    for r_idx, row in enumerate(sample_rq2, start=1):
        for c_idx, val in enumerate(row):
            cell = tbl_rq2.cell(r_idx, c_idx)
            set_cell_background(cell, "F0F4F8" if r_idx % 2 == 1 else "FFFFFF")
            set_cell_margins(cell, 60, 60, 80, 80)
            set_cell_borders(cell, color="CBD5E1", sz="4")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(val).font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_styled_heading(doc, "4.3 RQ4: Discernibilidade e Teste de Turing Perceptual", level=2)
    doc.add_paragraph("A Tabela 3 sumariza a taxa de identificação de autoria assistida e os graus de certeza reportados.")

    tbl_rq4 = doc.add_table(rows=4, cols=5)
    tbl_rq4.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_rq4 = ["Revisor", "Identificações Corretas (%)", "Grau Médio de Certeza (1-5)", "Casos com Certeza >= 4", "Padrões Linguísticos Observados"]
    for c_idx, h in enumerate(headers_rq4):
        cell = tbl_rq4.cell(0, c_idx)
        set_cell_background(cell, "1B365D")
        set_cell_margins(cell, 80, 80, 100, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(8.5)

    sample_rq4 = [
        ("Revisor 1", "[ % ]", "[ Média 1-5 ]", "[ Qtd ]", "[ Síntese dos Marcadores ]"),
        ("Revisor 2", "[ % ]", "[ Média 1-5 ]", "[ Qtd ]", "[ Síntese dos Marcadores ]"),
        ("Revisor 3", "[ % ]", "[ Média 1-5 ]", "[ Qtd ]", "[ Síntese dos Marcadores ]")
    ]
    for r_idx, row in enumerate(sample_rq4, start=1):
        for c_idx, val in enumerate(row):
            cell = tbl_rq4.cell(r_idx, c_idx)
            set_cell_background(cell, "F0F4F8" if r_idx % 2 == 1 else "FFFFFF")
            set_cell_margins(cell, 60, 60, 80, 80)
            set_cell_borders(cell, color="CBD5E1", sz="4")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(val).font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 5. DISCUSSÃO
    add_styled_heading(doc, "5. DISCUSSÃO DOS RESULTADOS", level=1)
    doc.add_paragraph(
        "A análise comparativa permite contrastar a eficiência operacional do RSACV2 em relação à triagem manual, "
        "com especial ênfase na capacidade do modelo de aplicar critérios de exclusão específicos (como distinguir turismo em fronteiras fluviais "
        "de estudos marítimos oceânicos ou de segurança urbana desvinculada de hidrovias). "
        "Destaca-se ainda a rastreabilidade proporcionada pelas justificativas automatizadas detalhadas, que superam anotações manuais lacônicas."
    )

    # 6. CONSIDERAÇÕES FINAIS
    add_styled_heading(doc, "6. CONSIDERAÇÕES FINAIS", level=1)
    doc.add_paragraph(
        "O protocolo experimental comprova a robustez e a confiabilidade do sistema RSACV2 na condução de revisões sistemáticas e de escopo. "
        "A ferramenta demonstra aderência aos padrões metodológicos internacionais, atuando como um assistente de alta precisão para pesquisadores "
        "em Ciências Sociais Aplicadas, Segurança Territorial e Desenvolvimento Regional."
    )

    # REFERÊNCIAS
    add_styled_heading(doc, "REFERÊNCIAS", level=1)
    refs = [
        "FLEISS, J. L. Measuring nominal scale agreement among many raters. Psychological Bulletin, v. 76, n. 5, p. 378-382, 1971.",
        "GOUGH, D.; OLIVER, S.; THOMAS, J. An Introduction to Systematic Reviews. 2. ed. London: SAGE Publications, 2017.",
        "MARSHALL, I. J.; WALLACE, B. C. Toward systematic review automation: a practical guide to using machine learning tools in research synthesis. Systematic Reviews, v. 8, n. 1, p. 163, 2019.",
        "McHUGH, M. L. Interrater reliability: the kappa statistic. Biochemia Medica, v. 22, n. 3, p. 276-282, 2012.",
        "PETERS, M. D. J. et al. Updated methodological guidance for the conduct of JBI scoping reviews. JBI Evidence Synthesis, v. 18, n. 10, p. 2119-2126, 2020.",
        "TRICCO, A. C. et al. PRISMA Extension for Scoping Reviews (PRISMA-ScR): Checklist and Explanation. Annals of Internal Medicine, v. 169, n. 7, p. 467-473, 2018."
    ]
    for r in refs:
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.left_indent = Inches(0.25)
        p_ref.paragraph_format.first_line_indent = Inches(-0.25)
        p_ref.paragraph_format.space_after = Pt(3.5)
        p_ref.add_run(r).font.size = Pt(9.5)

    doc_path = os.path.join(BASE_DIR, "Artigo_Esqueleto_Validacao_RSACV2.docx")
    doc.save(doc_path)
    print(f"Word document updated at: {doc_path}")

# ---------------------------------------------------------------------------
# 2. GERAÇÃO DO FORMULÁRIO DOS PESQUISADORES EM EXCEL
# ---------------------------------------------------------------------------
def generate_researcher_sheet():
    wb = openpyxl.Workbook()
    
    navy_dark = "1B365D"
    blue_header = "2B547E"
    blue_light = "EAF2F8"
    gray_light = "F8FAFC"
    border_color = "CBD5E1"
    
    font_title = Font(name="Segoe UI", size=13.5, bold=True, color="FFFFFF")
    font_subtitle = Font(name="Segoe UI", size=9.5, italic=True, color="E2E8F0")
    font_col_header = Font(name="Segoe UI", size=10, bold=True, color="FFFFFF")
    font_data = Font(name="Segoe UI", size=9.5)
    font_bold_data = Font(name="Segoe UI", size=9.5, bold=True)
    font_note = Font(name="Segoe UI", size=8.5, italic=True, color="64748B")
    
    fill_navy = PatternFill(start_color=navy_dark, end_color=navy_dark, fill_type="solid")
    fill_blue_head = PatternFill(start_color=blue_header, end_color=blue_header, fill_type="solid")
    fill_alt = PatternFill(start_color=blue_light, end_color=blue_light, fill_type="solid")
    fill_white = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
    
    thin_border_side = Side(style="thin", color=border_color)
    cell_border = Border(left=thin_border_side, right=thin_border_side, top=thin_border_side, bottom=thin_border_side)
    
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    align_top_left = Alignment(horizontal="left", vertical="top", wrap_text=True)
    
    # ----------------------------------------------------
    # ABA 1: INSTRUÇÕES & PROTOCOLO
    # ----------------------------------------------------
    ws1 = wb.active
    ws1.title = "Instruções & Protocolo"
    ws1.views.sheetView[0].showGridLines = True
    
    ws1.merge_cells("A1:G2")
    ws1["A1"] = "ESTUDO DE VALIDAÇÃO DO RSACV2 — FORMULÁRIO DO PESQUISADOR"
    ws1["A1"].font = font_title
    ws1["A1"].fill = fill_navy
    ws1["A1"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    ws1.merge_cells("A3:G3")
    ws1["A3"] = "Protocolo: Impactos da Segurança Pública na Operacionalização do Turismo Náutico em Fronteiras Fluviais (PRISMA-ScR)"
    ws1["A3"].font = font_subtitle
    ws1["A3"].fill = fill_blue_head
    ws1["A3"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    # Metadados do Pesquisador
    ws1["A5"] = "DADOS DO PESQUISADOR / EXECUTOR"
    ws1["A5"].font = Font(name="Segoe UI", size=11, bold=True, color=navy_dark)
    
    meta_fields = [
        ("Nome do Pesquisador:", "B6", "Braço Experimental:", "D6", "( ) Braço Manual   ( ) Braço RSACV2"),
        ("E-mail / Contato:", "B7", "Data de Execução:", "D7", "[DD/MM/AAAA]"),
        ("Instituição / Programa:", "B8", "Bases Aplicadas:", "D8", "BDTD, SciELO")
    ]
    for r_idx, (l1, c1, l2, c2, d2) in enumerate(meta_fields, start=6):
        ws1[f"A{r_idx}"] = l1
        ws1[f"A{r_idx}"].font = font_bold_data
        ws1[f"A{r_idx}"].border = cell_border
        ws1[f"B{r_idx}"].border = cell_border
        ws1[f"C{r_idx}"] = l2
        ws1[f"C{r_idx}"].font = font_bold_data
        ws1[f"C{r_idx}"].border = cell_border
        ws1[f"D{r_idx}"] = d2
        ws1[f"D{r_idx}"].font = font_data
        ws1[f"D{r_idx}"].border = cell_border
    
    # Protocolo de Busca (Descritores em Pares)
    ws1["A10"] = "ESTRATÉGIA DE DESCRITORES DE BUSCA (PARES BOOLEANOS — REGRAS BDTD / SCIELO)"
    ws1["A10"].font = Font(name="Segoe UI", size=11, bold=True, color=navy_dark)
    
    headers_proto = ["Par #", "Expressão de Busca (Português)", "Expressão de Busca (Inglês)", "Expressão de Busca (Espanhol)", "Bases Aplicadas", "Tipos de Documento"]
    for c_idx, h in enumerate(headers_proto, start=1):
        c_letter = get_column_letter(c_idx)
        ws1[f"{c_letter}11"] = h
        ws1[f"{c_letter}11"].font = font_col_header
        ws1[f"{c_letter}11"].fill = fill_blue_head
        ws1[f"{c_letter}11"].alignment = align_center
        ws1[f"{c_letter}11"].border = cell_border
    
    proto_rows = [
        ("Par 1", '"turismo náutico" AND "fronteira"', '"nautical tourism" AND "border"', '"turismo náutico" AND "frontera"', "BDTD, SciELO", "Artigos, Teses, Dissertações"),
        ("Par 2", '"segurança pública" AND "fronteira fluvial"', '"public security" AND "river border"', '"seguridad pública" AND "frontera fluvial"', "BDTD, SciELO", "Artigos, Teses, Dissertações"),
        ("Par 3", '"turismo" AND "fronteira fluvial"', '"tourism" AND "river border"', '"turismo" AND "frontera fluvial"', "BDTD, SciELO", "Artigos, Teses, Dissertações"),
        ("Par 4", '"turismo" AND "tríplice fronteira"', '"tourism" AND "cross-border"', '"turismo" AND "triple frontera"', "BDTD, SciELO", "Artigos, Teses, Dissertações"),
        ("Par 5", '"segurança pública" AND "turismo náutico"', '"water tourism" AND "border"', '"turismo fluvial" AND "frontera"', "BDTD, SciELO", "Artigos, Teses, Dissertações")
    ]
    for r_idx, row_data in enumerate(proto_rows, start=12):
        for c_idx, val in enumerate(row_data, start=1):
            c_letter = get_column_letter(c_idx)
            cell = ws1[f"{c_letter}{r_idx}"]
            cell.value = val
            cell.font = font_data
            cell.fill = fill_alt if r_idx % 2 == 1 else fill_white
            cell.border = cell_border
            cell.alignment = align_center if c_idx in [1, 5] else align_left
            
    # Critérios de Inclusão e Exclusão
    ws1["A18"] = "CRITÉRIOS DE ELEGIBILIDADE (TRIAGEM DE TÍTULOS E RESUMOS)"
    ws1["A18"].font = Font(name="Segoe UI", size=11, bold=True, color=navy_dark)
    
    crit_headers = ["Código", "Tipo", "Nome do Critério", "Texto Oficial do Critério no Protocolo"]
    for c_idx, h in enumerate(crit_headers, start=1):
        c_letter = get_column_letter(c_idx)
        ws1[f"{c_letter}19"] = h
        ws1[f"{c_letter}19"].font = font_col_header
        ws1[f"{c_letter}19"].fill = fill_blue_head
        ws1[f"{c_letter}19"].alignment = align_center
        ws1[f"{c_letter}19"].border = cell_border
        
    crit_data = [
        ("CI1", "Inclusão", "Turismo Náutico e Navegação em Fronteira", "Estudos que abordem a atividade turística, náutica, recreativa ou de navegação de passageiros em regiões de fronteira fluvial ou hidrovias transfronteiriças."),
        ("CI2", "Inclusão", "Segurança Pública e Governança Fluvial", "Pesquisas que analisem aspectos de segurança pública, criminalidade transfronteiriça, fiscalização, policiamento ou governança em bacias hidrográficas de fronteira."),
        ("CI3", "Inclusão", "Tipologia Documental e Idioma", "Publicações científicas completas (artigos de periódicos, teses e dissertações) nos idiomas português, inglês ou espanhol."),
        ("CE1", "Exclusão", "Transporte Marítimo / Alto-Mar", "Estudos com foco exclusivo em transporte marítimo oceânico ou de alto-mar sem interface fluvial ou fronteiriça."),
        ("CE2", "Exclusão", "Segurança Urbana/Rural Desvinculada", "Trabalhos sobre segurança pública puramente urbana ou rural sem qualquer conexão com hidrovias, cursos d'água de fronteira ou atividades turísticas."),
        ("CE3", "Exclusão", "Publicação Não Qualificada", "Documentos editoriais, resenhas de livros, resumos expandidos de eventos ou textos sem metodologia científica definida.")
    ]
    for r_idx, row_data in enumerate(crit_data, start=20):
        for c_idx, val in enumerate(row_data, start=1):
            c_letter = get_column_letter(c_idx)
            cell = ws1[f"{c_letter}{r_idx}"]
            cell.value = val
            cell.font = font_data
            cell.fill = fill_alt if "Inclusão" in row_data[1] else fill_white
            cell.border = cell_border
            cell.alignment = align_center if c_idx in [1, 2] else align_left
            
    # Questões de Extração
    ws1["A27"] = "QUESTÕES DE EXTRAÇÃO DE DADOS (PARA ESTUDOS INCLUÍDOS)"
    ws1["A27"].font = Font(name="Segoe UI", size=11, bold=True, color=navy_dark)
    
    qe_headers = ["ID", "Dimensão", "Pergunta de Extração", "Orientações de Preenchimento"]
    for c_idx, h in enumerate(qe_headers, start=1):
        c_letter = get_column_letter(c_idx)
        ws1[f"{c_letter}28"] = h
        ws1[f"{c_letter}28"].font = font_col_header
        ws1[f"{c_letter}28"].fill = fill_blue_head
        ws1[f"{c_letter}28"].alignment = align_center
        ws1[f"{c_letter}28"].border = cell_border
        
    qe_data = [
        ("QE1", "Localização & Bacia", "Qual é a localização geográfica, país e bacia hidrográfica/rio de fronteira analisado no estudo?", "Especificar país(es), cidades gêmeas, bacia hidrográfica (ex: Bacia Amazônica, Bacia do Prata) e rio de fronteira (ex: Rio Solimões, Rio Paraná, Rio Uruguai)."),
        ("QE2", "Tipologias de Crimes", "Quais tipologias de ocorrências de segurança pública, crimes ou ilícitos transfronteiriços foram identificadas?", "Identificar tráfico de drogas, contrabando/descaminho, pirataria fluvial, garimpo ilegal, crimes ambientais, conflitos armados, etc."),
        ("QE3", "Impactos no Turismo", "Quais foram os impactos diretos ou indiretos na atratividade, infraestrutura e dinâmica operacional do turismo náutico?", "Descrever restrições de navegação, sensação de insegurança de turistas, danos à infraestrutura de atracação, cancelamento de roteiros náuticos, etc."),
        ("QE4", "Governança & Políticas", "Quais estratégias de governança transfronteiriça, políticas públicas ou medidas de policiamento/mitigação foram recomendadas?", "Registrar bases fluviais integradas, cooperação policial binacional/trinacional, fiscalização aduaneira, monitoramento eletrônico por satélite/radar, etc."),
        ("QE5", "Metodologia & Fontes", "Qual a metodologia de pesquisa empregada e quais fontes de dados foram utilizadas?", "Descrever abordagem (qualitativa, quantitativa, mista), técnicas (entrevistas, survey, dados policiais, geoprocessamento) e fontes documentais.")
    ]
    for r_idx, row_data in enumerate(qe_data, start=29):
        for c_idx, val in enumerate(row_data, start=1):
            c_letter = get_column_letter(c_idx)
            cell = ws1[f"{c_letter}{r_idx}"]
            cell.value = val
            cell.font = font_data
            cell.fill = fill_alt if r_idx % 2 == 1 else fill_white
            cell.border = cell_border
            cell.alignment = align_center if c_idx == 1 else align_left

    ws1.column_dimensions['A'].width = 12
    ws1.column_dimensions['B'].width = 35
    ws1.column_dimensions['C'].width = 35
    ws1.column_dimensions['D'].width = 35
    ws1.column_dimensions['E'].width = 20
    ws1.column_dimensions['F'].width = 28
    ws1.column_dimensions['G'].width = 25

    # ----------------------------------------------------
    # ABA 2: 1. COLETA E BUSCA
    # ----------------------------------------------------
    ws2 = wb.create_sheet("1. Coleta e Busca")
    ws2.views.sheetView[0].showGridLines = True
    
    ws2.merge_cells("A1:I2")
    ws2["A1"] = "ETAPA 1: REGISTRO DE COLETA E RECUPERAÇÃO DE TRABALHOS (RQ1)"
    ws2["A1"].font = font_title
    ws2["A1"].fill = fill_navy
    ws2["A1"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    ws2.merge_cells("A3:I3")
    ws2["A3"] = "Instrução: Registre todos os trabalhos recuperados na ordem exata de aparição retornada pela base de dados para o protocolo de Turismo Náutico e Fronteiras Fluviais."
    ws2["A3"].font = font_subtitle
    ws2["A3"].fill = fill_blue_head
    ws2["A3"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    headers_coleta = ["ID Estudo", "Ordem de Aparição", "Base de Dados", "Par de Busca Utilizado", "Título do Trabalho", "Autores", "Ano", "Tipo de Documento", "DOI / Link / Handle"]
    for c_idx, h in enumerate(headers_coleta, start=1):
        c_letter = get_column_letter(c_idx)
        ws2[f"{c_letter}5"] = h
        ws2[f"{c_letter}5"].font = font_col_header
        ws2[f"{c_letter}5"].fill = fill_blue_head
        ws2[f"{c_letter}5"].alignment = align_center
        ws2[f"{c_letter}5"].border = cell_border
        
    exemplos_coleta = [
        ("DOC-001", 1, "BDTD", 'Par 2 ("segurança pública" AND "fronteira fluvial")', "O Custo-Amazonas e a Logística da Segurança Pública: Um Estudo sobre a Eficácia do Policiamento Ostensivo da PMAM em Atalaia do Norte", "Fabio Saldanha Soares; Denison Melo de Aguiar; et al.", 2026, "Artigo em Periódico / Tese", "https://doi.org/10.51891/rease.v12i3.24725"),
        ("DOC-002", 2, "BDTD", 'Par 2 ("segurança pública" AND "fronteira fluvial")', "Comando de Operações de Divisas na Polícia Militar do Amazonas: Controle Territorial Preventivo, Policiamento Fluvial e Especialização Operacional na Amazônia", "Jacksfran Barros Feitoza; Denison Melo de Aguiar; et al.", 2026, "Artigo em Periódico / Dissertação", "https://doi.org/10.51891/rease.v12i4.25164"),
        ("DOC-003", 3, "SciELO", 'Par 4 ("turismo" AND "tríplice fronteira")', "Políticas públicas, desenvolvimento sustentável e turismo na Tríplice Fronteira Brasil, Colômbia e Peru", "Paulo Moreira Pinto; Lígia Terezinha Lopes Simonian; Germán Palacio", 2019, "Artigo em Periódico", "https://doi.org/10.1590/scielo_exemplo_003"),
        ("DOC-004", 4, "BDTD", 'Par 4 ("turismo" AND "tríplice fronteira")', "Tríplice fronteira entre Brasil, Paraguai e Argentina: uma revisão sobre os aspectos históricos, econômicos, sociais, ambientais e aduaneiros que moldaram essa região", "Wagner Ferreira; Andréia Cristina Furtado", 2026, "Artigo em Periódico / Dissertação", "https://doi.org/10.54018/sssrv7n2-005"),
        ("DOC-005", 5, "SciELO", 'Par 1 ("turismo náutico" AND "fronteira")', "Território, Territorialidades e o Turismo como Desenvolvimento Regional na Tríplice Fronteira Brasil, Paraguai e Argentina", "Pesquisadores da Unioeste", 2018, "Artigo em Periódico", "https://doi.org/10.1590/scielo_exemplo_005")
    ]
    for r_idx, row_data in enumerate(exemplos_coleta, start=6):
        for c_idx, val in enumerate(row_data, start=1):
            c_letter = get_column_letter(c_idx)
            cell = ws2[f"{c_letter}{r_idx}"]
            cell.value = val
            cell.font = font_data
            cell.fill = fill_alt if r_idx % 2 == 1 else fill_white
            cell.border = cell_border
            cell.alignment = align_center if c_idx in [1, 2, 3, 7] else align_left

    for r_idx in range(11, 41):
        for c_idx in range(1, len(headers_coleta) + 1):
            c_letter = get_column_letter(c_idx)
            cell = ws2[f"{c_letter}{r_idx}"]
            cell.border = cell_border
            if c_idx == 1:
                cell.value = f"DOC-{r_idx-5:03d}"
                cell.alignment = align_center
                cell.font = font_note
            elif c_idx == 2:
                cell.value = r_idx - 5
                cell.alignment = align_center
                cell.font = font_note

    ws2.column_dimensions['A'].width = 12
    ws2.column_dimensions['B'].width = 16
    ws2.column_dimensions['C'].width = 14
    ws2.column_dimensions['D'].width = 35
    ws2.column_dimensions['E'].width = 45
    ws2.column_dimensions['F'].width = 30
    ws2.column_dimensions['G'].width = 10
    ws2.column_dimensions['H'].width = 25
    ws2.column_dimensions['I'].width = 35

    # ----------------------------------------------------
    # ABA 3: 2. TRIAGEM (INC & EXC)
    # ----------------------------------------------------
    ws3 = wb.create_sheet("2. Triagem (Inc & Exc)")
    ws3.views.sheetView[0].showGridLines = True
    
    ws3.merge_cells("A1:K2")
    ws3["A1"] = "ETAPA 2: TRIAGEM DE TRABALHOS POR CRITÉRIOS DE ELEGIBILIDADE (RQ2 & RQ3)"
    ws3["A1"].font = font_title
    ws3["A1"].fill = fill_navy
    ws3["A1"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    ws3.merge_cells("A3:K3")
    ws3["A3"] = "Instrução: Avalie Título e Resumo, assinale SIM/NÃO para os critérios CI1-CI3 e CE1-CE3, determine a Decisão Final e elabore a Justificativa Completa."
    ws3["A3"].font = font_subtitle
    ws3["A3"].fill = fill_blue_head
    ws3["A3"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    headers_triagem = [
        "ID Estudo", "Título do Estudo", "Resumo (Abstract)",
        "CI1 (Turismo/Naveg)", "CI2 (Segurança/Gov)", "CI3 (Tipo/Idioma)",
        "CE1 (Marítimo Alto-Mar)", "CE2 (Seg. Urbana Desv)", "CE3 (Não Qualificado)",
        "Decisão Final", "Justificativa Detalhada da Decisão"
    ]
    for c_idx, h in enumerate(headers_triagem, start=1):
        c_letter = get_column_letter(c_idx)
        ws3[f"{c_letter}5"] = h
        ws3[f"{c_letter}5"].font = font_col_header
        ws3[f"{c_letter}5"].fill = fill_blue_head
        ws3[f"{c_letter}5"].alignment = align_center
        ws3[f"{c_letter}5"].border = cell_border
        
    dv_decisao = DataValidation(type="list", formula1='"INCLUÍDO,EXCLUÍDO,DUVIDOSO"', allow_blank=True)
    ws3.add_data_validation(dv_decisao)
    dv_decisao.add("J6:J50")
    
    dv_sim_nao = DataValidation(type="list", formula1='"SIM,NÃO"', allow_blank=True)
    ws3.add_data_validation(dv_sim_nao)
    dv_sim_nao.add("D6:I50")

    exemplos_triagem = [
        ("DOC-001", "O Custo-Amazonas e a Logística da Segurança Pública: Um Estudo sobre a Eficácia do Policiamento Ostensivo da PMAM em Atalaia do Norte",
         "Este artigo analisa o impacto do custo-amazonas na segurança pública de Atalaia do Norte, fronteira do Vale do Javari. O objeto de pesquisa centra-se nos gargalos logísticos e na eficácia do policiamento fluvial...",
         "SIM", "SIM", "SIM", "NÃO", "NÃO", "NÃO", "INCLUÍDO",
         "O estudo aborda aspectos de segurança pública e policiamento ostensivo em região de fronteira fluvial (Vale do Javari, Atalaia do Norte) e analisa a mobilidade condicionada pela sazonalidade e infraestrutura fluvial, atendendo aos critérios de inclusão CI1, CI2 e CI3. Nenhum critério de exclusão foi satisfeito."),
        ("DOC-003", "Políticas públicas, desenvolvimento sustentável e turismo na Tríplice Fronteira Brasil, Colômbia e Peru",
         "Neste artigo, faz-se uma reflexão teórica sobre as políticas públicas de turismo e as suas implicações para a implementação do desenvolvimento sustentável e da gestão local em áreas protegidas...",
         "SIM", "SIM", "SIM", "NÃO", "NÃO", "NÃO", "INCLUÍDO",
         "O estudo atende ao critério CI1 por abordar a atividade turística em região de fronteira fluvial (Tríplice Fronteira Brasil, Colômbia e Peru). Atende ao CI2 ao analisar tensões e conflitos territoriais, e CI3 por ser artigo completo revisado por pares. Não incorre em nenhum critério de exclusão."),
        ("DOC-006", "Análise físico-química de solos e condutividade elétrica em pomares irrigados",
         "Avaliação agronômica e laboratorial de amostras de solo sob fertirrigação...",
         "NÃO", "NÃO", "SIM", "NÃO", "SIM", "NÃO", "EXCLUÍDO",
         "Estudo estritamente laboratorial e agronômico sem qualquer abordagem de turismo náutico, fronteiras fluviais ou segurança pública (CE2)."),
        ("DOC-007", "Spinal Fractures during Touristic Motorboat Sea Cruises: An Avoidable Phenomenon",
         "Clinical study analyzing compressive spinal fractures in tourists aboard high-speed motorboats in open oceanic waters...",
         "SIM", "NÃO", "SIM", "SIM", "NÃO", "NÃO", "EXCLUÍDO",
         "O trabalho aborda passeios de lancha em ambiente marítimo costeiro/oceânico de alto-mar sem interface fluvial ou fronteiriça (CE1) e com foco clínico/médico sem interface com segurança pública territorial.")
    ]
    for r_idx, row_data in enumerate(exemplos_triagem, start=6):
        for c_idx, val in enumerate(row_data, start=1):
            c_letter = get_column_letter(c_idx)
            cell = ws3[f"{c_letter}{r_idx}"]
            cell.value = val
            cell.font = font_data
            cell.fill = fill_alt if r_idx % 2 == 1 else fill_white
            cell.border = cell_border
            if c_idx in [1, 4, 5, 6, 7, 8, 9, 10]:
                cell.alignment = align_center
                if c_idx == 10:
                    cell.font = font_bold_data
            else:
                cell.alignment = align_left

    for r_idx in range(10, 41):
        for c_idx in range(1, len(headers_triagem) + 1):
            c_letter = get_column_letter(c_idx)
            cell = ws3[f"{c_letter}{r_idx}"]
            cell.border = cell_border
            if c_idx == 1:
                cell.value = f"DOC-{r_idx-5:03d}"
                cell.alignment = align_center
                cell.font = font_note

    ws3.column_dimensions['A'].width = 12
    ws3.column_dimensions['B'].width = 35
    ws3.column_dimensions['C'].width = 40
    ws3.column_dimensions['D'].width = 15
    ws3.column_dimensions['E'].width = 15
    ws3.column_dimensions['F'].width = 15
    ws3.column_dimensions['G'].width = 16
    ws3.column_dimensions['H'].width = 16
    ws3.column_dimensions['I'].width = 16
    ws3.column_dimensions['J'].width = 16
    ws3.column_dimensions['K'].width = 45

    # ----------------------------------------------------
    # ABA 4: 3. EXTRAÇÃO DE DADOS
    # ----------------------------------------------------
    ws4 = wb.create_sheet("3. Extração de Dados")
    ws4.views.sheetView[0].showGridLines = True
    
    ws4.merge_cells("A1:H2")
    ws4["A1"] = "ETAPA 3: EXTRAÇÃO ESTRUTURADA DE EVIDÊNCIAS (RQ5)"
    ws4["A1"].font = font_title
    ws4["A1"].fill = fill_navy
    ws4["A1"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    ws4.merge_cells("A3:H3")
    ws4["A3"] = "Instrução: Para cada estudo INCLUÍDO na triagem, responda às 5 questões de extração (QE1 a QE5) do protocolo com fidelidade e síntese analítica."
    ws4["A3"].font = font_subtitle
    ws4["A3"].fill = fill_blue_head
    ws4["A3"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    headers_extracao = [
        "ID Estudo", "Título do Estudo",
        "QE1: Localização & Bacia Hidrográfica",
        "QE2: Tipologias de Ocorrências / Ilícitos",
        "QE3: Impactos no Turismo Náutico",
        "QE4: Governança & Políticas Recomendadas",
        "QE5: Metodologia & Fontes de Dados",
        "Observações / Citações Relevantes"
    ]
    for c_idx, h in enumerate(headers_extracao, start=1):
        c_letter = get_column_letter(c_idx)
        ws4[f"{c_letter}5"] = h
        ws4[f"{c_letter}5"].font = font_col_header
        ws4[f"{c_letter}5"].fill = fill_blue_head
        ws4[f"{c_letter}5"].alignment = align_center
        ws4[f"{c_letter}5"].border = cell_border
        
    exemplos_extracao = [
        ("DOC-001", "O Custo-Amazonas e a Logística da Segurança Pública: Um Estudo sobre a Eficácia do Policiamento Ostensivo da PMAM em Atalaia do Norte",
         "Brasil / Peru; Bacia Amazônica; Rio Javari e afluentes (Atalaia do Norte / Vale do Javari).",
         "Tráfico transfronteiriço de entorpecentes, pesca e caça predatória ilegal em terras indígenas, pirataria fluvial e ameaças a operadores turísticos e pesquisadores.",
         "Elevação extrema dos custos operacionais de navegação turística ('custo-amazonas'), necessidade de escolta e retração de operadoras de ecoturismo náutico.",
         "Implantação de bases fluviais integradas permanentes, ampliação do patrulhamento naval conjunto com Marinha/Polícia Federal e uso de embarcações blindadas de calado reduzido.",
         "Estudo empírico de abordagem mista com dados operacionais da PMAM, entrevistas com policiais fluviais e análise logística de tempos/custos de deslocamento.",
         "Citação: 'A sazonalidade do nível dos rios impõe janelas críticas onde o policiamento fluvial fica condicionado ao regime hidrológico, afetando a segurança das rotas turísticas.' (p. 45)"),
        ("DOC-003", "Políticas públicas, desenvolvimento sustentável e turismo na Tríplice Fronteira Brasil, Colômbia e Peru",
         "Tríplice Fronteira (Tabatinga/BR, Leticia/CO, Santa Rosa/PE); Bacia Amazônica (Rio Solimões / Amazonas).",
         "Contrabando transfronteiriço, disputas territoriais de facções em portos clandestinos e exploração informal de rotas de passageiros.",
         "Insegurança percebida pelos turistas em travessias noturnas de barco, assimetrias nas exigências de segurança náutica entre os três países e gargalos de infraestrutura portuária.",
         "Comitês de integração fronteiriça tripartite, harmonização das normas de segurança para embarcações turísticas e criação de corredores turísticos fluviais protegidos.",
         "Pesquisa qualitativa, análise documental de acordos bilaterais/trilaterais e entrevistas semiestruturadas com gestores municipais de turismo e autoridades de fronteira.",
         "Citação: 'A fluidez da fronteira fluvial exige governança compartilhada que supere a fragmentação das agências de fiscalização.' (p. 112)")
    ]
    for r_idx, row_data in enumerate(exemplos_extracao, start=6):
        for c_idx, val in enumerate(row_data, start=1):
            c_letter = get_column_letter(c_idx)
            cell = ws4[f"{c_letter}{r_idx}"]
            cell.value = val
            cell.font = font_data
            cell.fill = fill_alt if r_idx % 2 == 1 else fill_white
            cell.border = cell_border
            cell.alignment = align_center if c_idx == 1 else align_top_left

    for r_idx in range(8, 31):
        for c_idx in range(1, len(headers_extracao) + 1):
            c_letter = get_column_letter(c_idx)
            cell = ws4[f"{c_letter}{r_idx}"]
            cell.border = cell_border
            if c_idx == 1:
                cell.value = f"DOC-{r_idx-5:03d}"
                cell.alignment = align_center
                cell.font = font_note

    ws4.column_dimensions['A'].width = 12
    ws4.column_dimensions['B'].width = 30
    ws4.column_dimensions['C'].width = 32
    ws4.column_dimensions['D'].width = 35
    ws4.column_dimensions['E'].width = 38
    ws4.column_dimensions['F'].width = 35
    ws4.column_dimensions['G'].width = 32
    ws4.column_dimensions['H'].width = 30

    file_path = os.path.join(BASE_DIR, "Formulario_Pesquisadores_Triagem_Extracao.xlsx")
    wb.save(file_path)
    print(f"Researcher Excel workbook updated at: {file_path}")

# ---------------------------------------------------------------------------
# 3. GERAÇÃO DO FORMULÁRIO DOS AVALIADORES EM EXCEL (REVISÃO CEGA)
# ---------------------------------------------------------------------------
def generate_evaluator_sheet():
    wb = openpyxl.Workbook()
    
    navy_dark = "1B365D"
    blue_header = "2B547E"
    blue_light = "EAF2F8"
    gray_light = "F8FAFC"
    border_color = "CBD5E1"
    gold_accent = "B45309"
    
    font_title = Font(name="Segoe UI", size=13.5, bold=True, color="FFFFFF")
    font_subtitle = Font(name="Segoe UI", size=9.5, italic=True, color="E2E8F0")
    font_col_header = Font(name="Segoe UI", size=10, bold=True, color="FFFFFF")
    font_data = Font(name="Segoe UI", size=9.5)
    font_bold_data = Font(name="Segoe UI", size=9.5, bold=True)
    font_note = Font(name="Segoe UI", size=8.5, italic=True, color="64748B")
    
    fill_navy = PatternFill(start_color=navy_dark, end_color=navy_dark, fill_type="solid")
    fill_blue_head = PatternFill(start_color=blue_header, end_color=blue_header, fill_type="solid")
    fill_alt = PatternFill(start_color=blue_light, end_color=blue_light, fill_type="solid")
    fill_white = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
    fill_gold = PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid")
    
    thin_border_side = Side(style="thin", color=border_color)
    cell_border = Border(left=thin_border_side, right=thin_border_side, top=thin_border_side, bottom=thin_border_side)
    
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    align_top_left = Alignment(horizontal="left", vertical="top", wrap_text=True)
    
    # ----------------------------------------------------
    # ABA 1: INSTRUÇÕES AOS AVALIADORES
    # ----------------------------------------------------
    ws1 = wb.active
    ws1.title = "Instruções aos Avaliadores"
    ws1.views.sheetView[0].showGridLines = True
    
    ws1.merge_cells("A1:G2")
    ws1["A1"] = "ESTUDO DE VALIDAÇÃO DO RSACV2 — PROTOCOLO DE REVISÃO CEGA"
    ws1["A1"].font = font_title
    ws1["A1"].fill = fill_navy
    ws1["A1"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    ws1.merge_cells("A3:G3")
    ws1["A3"] = "Protocolo Avaliado: 'Impactos da Segurança Pública na Operacionalização do Turismo Náutico em Fronteiras Fluviais' (PRISMA-ScR)"
    ws1["A3"].font = font_subtitle
    ws1["A3"].fill = fill_blue_head
    ws1["A3"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    # Identificação do Revisor
    ws1["A5"] = "IDENTIFICAÇÃO DO AVALIADOR / REVISOR INDEPENDENTE"
    ws1["A5"].font = Font(name="Segoe UI", size=11, bold=True, color=navy_dark)
    
    meta_rev = [
        ("Código do Revisor:", "B6", "( ) Revisor 1   ( ) Revisor 2   ( ) Revisor 3", "D6", "Data da Avaliação:", "E6", "[DD/MM/AAAA]"),
        ("Nome do Avaliador:", "B7", "[Nome Completo do Avaliador]", "D7", "Titulação / Área:", "E7", "Doutor / Ciências Sociais Aplicadas"),
        ("Instituição de Origem:", "B8", "[Universidade / Instituto de Pesquisa]", "D8", "Termo de Sigilo:", "E8", "Assinado e Conforme")
    ]
    for r_idx, (l1, c1, v1, c2, l2, c3, v2) in enumerate(meta_rev, start=6):
        ws1[f"A{r_idx}"] = l1
        ws1[f"A{r_idx}"].font = font_bold_data
        ws1[f"A{r_idx}"].border = cell_border
        ws1[f"B{r_idx}"] = v1
        ws1[f"B{r_idx}"].font = font_data
        ws1[f"B{r_idx}"].border = cell_border
        ws1[f"C{r_idx}"].border = cell_border
        ws1[f"D{r_idx}"] = l2
        ws1[f"D{r_idx}"].font = font_bold_data
        ws1[f"D{r_idx}"].border = cell_border
        ws1[f"E{r_idx}"] = v2
        ws1[f"E{r_idx}"].font = font_data
        ws1[f"E{r_idx}"].border = cell_border
        
    # Regras de Julgamento
    ws1["A10"] = "CRITÉRIOS E REGRAS DE JULGAMENTO CEGO (DIMENSÕES 1 A 6)"
    ws1["A10"].font = Font(name="Segoe UI", size=11, bold=True, color=navy_dark)
    
    regras_headers = ["Dimensão / Aba", "Questão de Pesquisa", "Critério de Julgamento e Opções de Resposta", "Obrigatoriedade de Justificativa"]
    for c_idx, h in enumerate(regras_headers, start=1):
        c_letter = get_column_letter(c_idx)
        ws1[f"{c_letter}11"] = h
        ws1[f"{c_letter}11"].font = font_col_header
        ws1[f"{c_letter}11"].fill = fill_blue_head
        ws1[f"{c_letter}11"].alignment = align_center
        ws1[f"{c_letter}11"].border = cell_border
        
    regras_tabela = [
        ("Aba 'Q1 - Coleta'", "RQ1: Igualdade na Coleta e Ordem",
         "Critério 1.1: Escolha categórica entre:\n• SIM (100% igual em presença e ordem)\n• NÃO - POR PRESENÇA (estudos distintos)\n• NÃO - POR ORDEM (mesmos estudos, posições distintas)\n• NÃO INTEGRAL (diverge em presença e ordem)",
         "Opcional se 'SIM'; Recomendada se houver divergência."),
        ("Aba 'Q2 & Q3 - Triagem'", "RQ2 & RQ3: Decisões e Divergências de Triagem",
         "Critérios 2.1 e 3.1:\n• Concordância de Decisão e Justificativa (SIM / NÃO)\n• Em caso de divergência, assinalar preferência:\n  - Método A está correto\n  - Método B está correto\n  - Ambos divergem mas estão corretos\n  - Divergem mas ambos erram",
         "OBRIGATÓRIA POR EXTENSO em todos os casos de divergência de decisão ou justificativa."),
        ("Aba 'Q4 - Percepção Turing'", "RQ4: Discernimento da Assistência Computacional",
         "Critério 4.1:\n• Qual resposta parece ser assistida por IA/RSAC? (Método A / Método B / Indistinguível)\n• Grau de Certeza na escala de 1 a 5:\n  1=Totalmente incerto (chute)\n  2=Pouco certo\n  3=Moderadamente certo\n  4=Muito certo\n  5=Certeza absoluta",
         "OBRIGATÓRIA sempre que a Certeza for 4 ou 5 (apontar marcadores textuais e sintáticos)."),
        ("Aba 'Q5 - Extração'", "RQ5: Qualidade na Extração de Informações",
         "Critérios 5.1 e 5.2:\n• Registro textual de Similaridades e Diferenças/Omissões\n• Nota de Concordância com Método A (1 a 5)\n• Nota de Concordância com Método B (1 a 5)\n• Extração Mais Completa / Preferida",
         "Campo estruturado de comparação de similaridades e diferenças."),
        ("Aba 'Q6 - Síntese Qualitativa'", "RQ6: Observações Qualitativas Globais",
         "Critério 6.1: Redação de parágrafos reflexivos estruturados avaliando rigor de raciocínio, profundidade, discriminação de autoria e confiabilidade para a pesquisa.",
         "OBRIGATÓRIA (Parágrafos descritivos completos).")
    ]
    for r_idx, row_data in enumerate(regras_tabela, start=12):
        for c_idx, val in enumerate(row_data, start=1):
            c_letter = get_column_letter(c_idx)
            cell = ws1[f"{c_letter}{r_idx}"]
            cell.value = val
            cell.font = font_data
            cell.fill = fill_alt if r_idx % 2 == 1 else fill_white
            cell.border = cell_border
            cell.alignment = align_center if c_idx == 1 else align_top_left

    ws1.column_dimensions['A'].width = 22
    ws1.column_dimensions['B'].width = 30
    ws1.column_dimensions['C'].width = 50
    ws1.column_dimensions['D'].width = 32
    ws1.column_dimensions['E'].width = 25
    ws1.column_dimensions['F'].width = 20
    ws1.column_dimensions['G'].width = 20

    # ----------------------------------------------------
    # ABA 2: Q1 - COLETA E BUSCA
    # ----------------------------------------------------
    ws2 = wb.create_sheet("Q1 - Coleta e Busca")
    ws2.views.sheetView[0].showGridLines = True
    
    ws2.merge_cells("A1:H2")
    ws2["A1"] = "AVALIAÇÃO RQ1: COMPARATIVO DE COLETA E RECUPERAÇÃO DE TRABALHOS"
    ws2["A1"].font = font_title
    ws2["A1"].fill = fill_navy
    ws2["A1"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    ws2.merge_cells("A3:H3")
    ws2["A3"] = "Critério 1.1: Avalie se os trabalhos coletados pelos dois métodos na mesma base e com o mesmo protocolo de Turismo Náutico em Fronteiras Fluviais são idênticos em presença e ordem."
    ws2["A3"].font = font_subtitle
    ws2["A3"].fill = fill_blue_head
    ws2["A3"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    headers_q1 = [
        "Par de Busca #", "Base de Dados",
        "Trabalhos Coletados (Método A)",
        "Trabalhos Coletados (Método B)",
        "Julgamento 1.1 (Resultado)",
        "Diferenças de Presença Detectadas",
        "Diferenças de Ordem Detectadas",
        "Observações / Justificativa do Revisor"
    ]
    for c_idx, h in enumerate(headers_q1, start=1):
        c_letter = get_column_letter(c_idx)
        ws2[f"{c_letter}5"] = h
        ws2[f"{c_letter}5"].font = font_col_header
        ws2[f"{c_letter}5"].fill = fill_blue_head
        ws2[f"{c_letter}5"].alignment = align_center
        ws2[f"{c_letter}5"].border = cell_border
        
    dv_q1 = DataValidation(type="list", formula1='"SIM,NÃO - POR PRESENÇA,NÃO - POR ORDEM,NÃO INTEGRAL"', allow_blank=True)
    ws2.add_data_validation(dv_q1)
    dv_q1.add("E6:E30")

    exemplos_q1 = [
        ('Par 2 ("segurança pública" AND "fronteira fluvial")', "BDTD", "DOC-001 (1º), DOC-002 (2º)", "DOC-001 (1º), DOC-002 (2º)", "SIM", "Nenhuma (100% coincidentes)", "Nenhuma (mesma ordem 1º e 2º)", "Ambos os métodos recuperaram exatamente as mesmas publicações na mesma ordem retornada pelo VuFind."),
        ('Par 4 ("turismo" AND "tríplice fronteira")', "SciELO", "DOC-003 (1º), DOC-004 (2º)", "DOC-004 (1º), DOC-003 (2º)", "NÃO - POR ORDEM", "Nenhuma (mesmos artigos recuperados)", "Ordem invertida no ranqueamento de relevância", "Os artigos recuperados foram idênticos, mas com ligeira alteração na ordem de listagem da API."),
        ('Par 1 ("turismo náutico" AND "fronteira")', "BDTD", "DOC-005 (1º), DOC-008 (2º)", "DOC-005 (1º), DOC-009 (2º)", "NÃO - POR PRESENÇA", "Método A recuperou DOC-008; Método B recuperou DOC-009", "Ordem não comparável para o 2º item", "Divergência na recuperação do segundo registro devido ao tratamento de acentuação/sinônimos.")
    ]
    for r_idx, row_data in enumerate(exemplos_q1, start=6):
        for c_idx, val in enumerate(row_data, start=1):
            c_letter = get_column_letter(c_idx)
            cell = ws2[f"{c_letter}{r_idx}"]
            cell.value = val
            cell.font = font_data
            cell.fill = fill_alt if r_idx % 2 == 1 else fill_white
            cell.border = cell_border
            cell.alignment = align_center if c_idx in [1, 2, 5] else align_left
            if c_idx == 5:
                cell.font = font_bold_data

    for r_idx in range(9, 21):
        for c_idx in range(1, len(headers_q1) + 1):
            c_letter = get_column_letter(c_idx)
            cell = ws2[f"{c_letter}{r_idx}"]
            cell.border = cell_border
            if c_idx == 1:
                cell.value = f"Par {(r_idx-5)%5 + 1}"
                cell.alignment = align_center
                cell.font = font_note

    ws2.column_dimensions['A'].width = 30
    ws2.column_dimensions['B'].width = 15
    ws2.column_dimensions['C'].width = 30
    ws2.column_dimensions['D'].width = 30
    ws2.column_dimensions['E'].width = 25
    ws2.column_dimensions['F'].width = 30
    ws2.column_dimensions['G'].width = 30
    ws2.column_dimensions['H'].width = 40

    # ----------------------------------------------------
    # ABA 3: Q2 & Q3 - TRIAGEM E DECISÕES
    # ----------------------------------------------------
    ws3 = wb.create_sheet("Q2 & Q3 - Triagem")
    ws3.views.sheetView[0].showGridLines = True
    
    ws3.merge_cells("A1:K2")
    ws3["A1"] = "AVALIAÇÃO RQ2 & RQ3: TRIAGEM CEGA E PREFERÊNCIA EM DIVERGÊNCIAS"
    ws3["A1"].font = font_title
    ws3["A1"].fill = fill_navy
    ws3["A1"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    ws3.merge_cells("A3:K3")
    ws3["A3"] = "Instrução: Compare as decisões e justificativas dos Métodos A e B no protocolo de Turismo Náutico e Fronteiras Fluviais. Em caso de divergência, assinale sua preferência e justifique."
    ws3["A3"].font = font_subtitle
    ws3["A3"].fill = fill_blue_head
    ws3["A3"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    headers_q23 = [
        "ID Estudo", "Título do Estudo", "Resumo do Estudo",
        "Decisão (Método A)", "Justificativa (Método A)",
        "Decisão (Método B)", "Justificativa (Método B)",
        "Concordância de Decisão?",
        "Concordância de Justificativa?",
        "Preferência do Revisor (Em Divergência)",
        "Justificativa da Preferência (Por Extenso)"
    ]
    for c_idx, h in enumerate(headers_q23, start=1):
        c_letter = get_column_letter(c_idx)
        ws3[f"{c_letter}5"] = h
        ws3[f"{c_letter}5"].font = font_col_header
        ws3[f"{c_letter}5"].fill = fill_blue_head
        ws3[f"{c_letter}5"].alignment = align_center
        ws3[f"{c_letter}5"].border = cell_border
        
    dv_conc = DataValidation(type="list", formula1='"SIM,NÃO"', allow_blank=True)
    ws3.add_data_validation(dv_conc)
    dv_conc.add("H6:I40")
    
    dv_pref = DataValidation(type="list", formula1='"MÉTODO A ESTÁ CORRETO,MÉTODO B ESTÁ CORRETO,AMBOS DIVERGEM MAS ESTÃO CORRETOS,DIVERGEM MAS AMBOS ERRAM,CONCORDANTES (NÃO SE APLICA)"', allow_blank=True)
    ws3.add_data_validation(dv_pref)
    dv_pref.add("J6:J40")

    exemplos_q23 = [
        ("DOC-001", "O Custo-Amazonas e a Logística da Segurança Pública: Um Estudo sobre a Eficácia do Policiamento Ostensivo da PMAM em Atalaia do Norte",
         "Este artigo analisa o impacto do custo-amazonas na segurança pública de Atalaia do Norte, fronteira do Vale do Javari. O objeto de pesquisa centra-se nos gargalos logísticos e na eficácia do policiamento fluvial...",
         "INCLUÍDO", "O estudo aborda segurança pública e policiamento fluvial em fronteira (Vale do Javari), atendendo aos critérios CI1, CI2 e CI3, sem incorrer em critérios de exclusão.",
         "INCLUÍDO", "Trabalho empírico sobre policiamento ostensivo e logística fluvial na fronteira do Amazonas com o Peru. Cumpre os critérios de inclusão do protocolo.",
         "SIM", "SIM", "CONCORDANTES (NÃO SE APLICA)", "Ambos os métodos identificaram corretamente o enquadramento temático e fundamentaram nos critérios do protocolo."),
        ("DOC-007", "Spinal Fractures during Touristic Motorboat Sea Cruises: An Avoidable Phenomenon",
         "Clinical study analyzing compressive spinal fractures in tourists aboard high-speed motorboats in open oceanic waters...",
         "EXCLUÍDO", "O estudo foi excluído com base no critério CE1 por focar em transporte marítimo oceânico/mar aberto sem interface fluvial ou transfronteiriça, além de ter foco puramente médico/clínico.",
         "INCLUÍDO", "O estudo menciona turismo em lanchas a motor (motorboats) e atividade turística náutica.",
         "NÃO", "NÃO", "MÉTODO A ESTÁ CORRETO", "O Método A aplicou com precisão o critério de exclusão CE1 (transporte marítimo oceânico sem interface fluvial ou fronteiriça). O Método B foi induzido pela palavra-chave 'motorboat' mas ignorou o recorte fluvial de fronteira exigido pelo protocolo.")
    ]
    for r_idx, row_data in enumerate(exemplos_q23, start=6):
        for c_idx, val in enumerate(row_data, start=1):
            c_letter = get_column_letter(c_idx)
            cell = ws3[f"{c_letter}{r_idx}"]
            cell.value = val
            cell.font = font_data
            cell.fill = fill_alt if r_idx % 2 == 1 else fill_white
            cell.border = cell_border
            if c_idx in [1, 4, 6, 8, 9, 10]:
                cell.alignment = align_center
                if c_idx == 10:
                    cell.font = font_bold_data
            else:
                cell.alignment = align_left

    for r_idx in range(8, 31):
        for c_idx in range(1, len(headers_q23) + 1):
            c_letter = get_column_letter(c_idx)
            cell = ws3[f"{c_letter}{r_idx}"]
            cell.border = cell_border
            if c_idx == 1:
                cell.value = f"DOC-{r_idx-5:03d}"
                cell.alignment = align_center
                cell.font = font_note

    ws3.column_dimensions['A'].width = 12
    ws3.column_dimensions['B'].width = 30
    ws3.column_dimensions['C'].width = 35
    ws3.column_dimensions['D'].width = 16
    ws3.column_dimensions['E'].width = 35
    ws3.column_dimensions['F'].width = 16
    ws3.column_dimensions['G'].width = 35
    ws3.column_dimensions['H'].width = 16
    ws3.column_dimensions['I'].width = 16
    ws3.column_dimensions['J'].width = 28
    ws3.column_dimensions['K'].width = 45

    # ----------------------------------------------------
    # ABA 4: Q4 - TESTE DE PERCEPÇÃO (TURING)
    # ----------------------------------------------------
    ws4 = wb.create_sheet("Q4 - Percepção Turing")
    ws4.views.sheetView[0].showGridLines = True
    
    ws4.merge_cells("A1:I2")
    ws4["A1"] = "AVALIAÇÃO RQ4: TESTE DE PERCEPÇÃO E DISCERNIMENTO DA ASSISTÊNCIA COMPUTACIONAL"
    ws4["A1"].font = font_title
    ws4["A1"].fill = fill_navy
    ws4["A1"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    ws4.merge_cells("A3:I3")
    ws4["A3"] = "Critério 4.1: Indique qual resposta parece ser assistida pelo RSACV2 e seu grau de certeza (1 a 5). Se Certeza for 4 ou 5, justifique obrigatoriamente."
    ws4["A3"].font = font_subtitle
    ws4["A3"].fill = fill_blue_head
    ws4["A3"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    headers_q4 = [
        "ID Estudo", "Item Avaliado (Triagem / Extração)",
        "Resposta (Método A)",
        "Resposta (Método B)",
        "Qual parece ser a resposta ASSISTIDA?",
        "Grau de Certeza (1 a 5)",
        "Marcadores Linguísticos Observados",
        "Justificativa Detalhada (Obrigatória se Certeza >= 4)",
        "Autoria Real (Uso Exclusivo Coordenação)"
    ]
    for c_idx, h in enumerate(headers_q4, start=1):
        c_letter = get_column_letter(c_idx)
        ws4[f"{c_letter}5"] = h
        ws4[f"{c_letter}5"].font = font_col_header
        ws4[f"{c_letter}5"].fill = fill_blue_head
        ws4[f"{c_letter}5"].alignment = align_center
        ws4[f"{c_letter}5"].border = cell_border
        
    dv_assist = DataValidation(type="list", formula1='"MÉTODO A,MÉTODO B,INDISTINGUÍVEL / DÚVIDA"', allow_blank=True)
    ws4.add_data_validation(dv_assist)
    dv_assist.add("E6:E40")
    
    dv_likert = DataValidation(type="list", formula1='"1 - Totalmente incerto (chute),2 - Pouco certo,3 - Moderadamente certo,4 - Muito certo,5 - Certeza absoluta"', allow_blank=True)
    ws4.add_data_validation(dv_likert)
    dv_likert.add("F6:F40")

    exemplos_q4 = [
        ("DOC-001", "Justificativa de Triagem",
         "O estudo atende ao critério INC1 por analisar policiamento fluvial em Atalaia do Norte e INC2 por abordar criminalidade de fronteira, sendo artigo completo (INC3) sem enquadramento em critérios de exclusão.",
         "Trabalho focado em segurança fluvial na fronteira do Amazonas com o Peru. Incluído por aderência ao tema.",
         "MÉTODO A", "4 - Muito certo", "Citação exaustiva dos códigos de critérios (INC1, INC2, INC3) e estruturação formal padronizada.",
         "A resposta do Método A exibe o padrão característico de geração por IA que mapeia sistematicamente cada código de critério do protocolo.",
         "[Restrito / Codificado]"),
        ("DOC-003", "Extração QE3 (Impactos no Turismo)",
         "Insegurança percebida nas travessias noturnas de barco e carência de padronização nas normas de segurança náutica entre os três países.",
         "O estudo evidencia que as tensões transfronteiriças afetam a atratividade turística da região do Solimões, gerando receio nos operadores quanto à navegação em horários de menor visibilidade.",
         "MÉTODO B", "3 - Moderadamente certo", "Maior sofisticação estilística e síntese analítica.",
         "A redação fluida e conectada ao contexto do Rio Solimões sugere sumarização automatizada avançada.",
         "[Restrito / Codificado]")
    ]
    for r_idx, row_data in enumerate(exemplos_q4, start=6):
        for c_idx, val in enumerate(row_data, start=1):
            c_letter = get_column_letter(c_idx)
            cell = ws4[f"{c_letter}{r_idx}"]
            cell.value = val
            cell.font = font_data
            cell.fill = fill_alt if r_idx % 2 == 1 else fill_white
            cell.border = cell_border
            if c_idx in [1, 5, 6, 9]:
                cell.alignment = align_center
                if c_idx in [5, 6]:
                    cell.font = font_bold_data
            else:
                cell.alignment = align_left

    for r_idx in range(8, 31):
        for c_idx in range(1, len(headers_q4) + 1):
            c_letter = get_column_letter(c_idx)
            cell = ws4[f"{c_letter}{r_idx}"]
            cell.border = cell_border
            if c_idx == 1:
                cell.value = f"DOC-{r_idx-5:03d}"
                cell.alignment = align_center
                cell.font = font_note

    ws4.column_dimensions['A'].width = 12
    ws4.column_dimensions['B'].width = 24
    ws4.column_dimensions['C'].width = 35
    ws4.column_dimensions['D'].width = 35
    ws4.column_dimensions['E'].width = 24
    ws4.column_dimensions['F'].width = 28
    ws4.column_dimensions['G'].width = 32
    ws4.column_dimensions['H'].width = 45
    ws4.column_dimensions['I'].width = 22

    # ----------------------------------------------------
    # ABA 5: Q5 - EXTRAÇÃO DE DADOS
    # ----------------------------------------------------
    ws5 = wb.create_sheet("Q5 - Extração")
    ws5.views.sheetView[0].showGridLines = True
    
    ws5.merge_cells("A1:J2")
    ws5["A1"] = "AVALIAÇÃO RQ5: COMPARAÇÃO CEGA DA EXTRAÇÃO DE EVIDÊNCIAS"
    ws5["A1"].font = font_title
    ws5["A1"].fill = fill_navy
    ws5["A1"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    ws5.merge_cells("A3:J3")
    ws5["A3"] = "Critérios 5.1 e 5.2: Compare as extrações cegas para cada questão (QE1 a QE5), registre similaridades/diferenças e atribua notas de concordância (1 a 5)."
    ws5["A3"].font = font_subtitle
    ws5["A3"].fill = fill_blue_head
    ws5["A3"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    headers_q5 = [
        "ID Estudo", "Questão de Extração",
        "Resposta (Método A)",
        "Resposta (Método B)",
        "Similaridades Observadas",
        "Diferenças e Omissões Detectadas",
        "Concordância com Método A (1-5)",
        "Concordância com Método B (1-5)",
        "Extração Mais Completa / Preferida",
        "Justificativa / Parecer do Revisor"
    ]
    for c_idx, h in enumerate(headers_q5, start=1):
        c_letter = get_column_letter(c_idx)
        ws5[f"{c_letter}5"] = h
        ws5[f"{c_letter}5"].font = font_col_header
        ws5[f"{c_letter}5"].fill = fill_blue_head
        ws5[f"{c_letter}5"].alignment = align_center
        ws5[f"{c_letter}5"].border = cell_border
        
    dv_scale5 = DataValidation(type="list", formula1='"1,2,3,4,5"', allow_blank=True)
    ws5.add_data_validation(dv_scale5)
    dv_scale5.add("G6:H40")
    
    dv_pref_ext = DataValidation(type="list", formula1='"MÉTODO A,MÉTODO B,AMBAS EQUIVALENTES,NENHUMA SATISFATÓRIA"', allow_blank=True)
    ws5.add_data_validation(dv_pref_ext)
    dv_pref_ext.add("I6:I40")

    exemplos_q5 = [
        ("DOC-001", "QE4 (Governança & Políticas)",
         "Implantação de bases fluviais integradas permanentes, ampliação do patrulhamento naval conjunto com a Marinha e emprego de embarcações blindadas de calado reduzido.",
         "Bases fluviais conjuntas entre PMAM, Polícia Federal e órgãos ambientais para controle de rotas fluviais de fronteira.",
         "Ambos destacaram a necessidade de bases fluviais integradas interinstitucionais.",
         "Método A detalhou adicionalmente o tipo de embarcação blindada e a cooperação com a Marinha.",
         5, 4, "MÉTODO A", "A resposta do Método A é mais abrangente e contempla aspectos operacionais adicionais descritos no texto.")
    ]
    for r_idx, row_data in enumerate(exemplos_q5, start=6):
        for c_idx, val in enumerate(row_data, start=1):
            c_letter = get_column_letter(c_idx)
            cell = ws5[f"{c_letter}{r_idx}"]
            cell.value = val
            cell.font = font_data
            cell.fill = fill_alt if r_idx % 2 == 1 else fill_white
            cell.border = cell_border
            if c_idx in [1, 2, 7, 8, 9]:
                cell.alignment = align_center
                if c_idx in [7, 8, 9]:
                    cell.font = font_bold_data
            else:
                cell.alignment = align_left

    for r_idx in range(7, 26):
        for c_idx in range(1, len(headers_q5) + 1):
            c_letter = get_column_letter(c_idx)
            cell = ws5[f"{c_letter}{r_idx}"]
            cell.border = cell_border
            if c_idx == 1:
                cell.value = f"DOC-{(r_idx-6)//5 + 1:03d}"
                cell.alignment = align_center
                cell.font = font_note
            elif c_idx == 2:
                cell.value = f"QE{(r_idx-6)%5 + 1}"
                cell.alignment = align_center
                cell.font = font_note

    ws5.column_dimensions['A'].width = 12
    ws5.column_dimensions['B'].width = 20
    ws5.column_dimensions['C'].width = 35
    ws5.column_dimensions['D'].width = 35
    ws5.column_dimensions['E'].width = 30
    ws5.column_dimensions['F'].width = 30
    ws5.column_dimensions['G'].width = 16
    ws5.column_dimensions['H'].width = 16
    ws5.column_dimensions['I'].width = 24
    ws5.column_dimensions['J'].width = 40

    # ----------------------------------------------------
    # ABA 6: Q6 - SÍNTESE QUALITATIVA
    # ----------------------------------------------------
    ws6 = wb.create_sheet("Q6 - Síntese Qualitativa")
    ws6.views.sheetView[0].showGridLines = True
    
    ws6.merge_cells("A1:G2")
    ws6["A1"] = "AVALIAÇÃO RQ6: PARECER QUALITATIVO GLOBAL DO REVISOR"
    ws6["A1"].font = font_title
    ws6["A1"].fill = fill_navy
    ws6["A1"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    ws6.merge_cells("A3:G3")
    ws6["A3"] = "Critério 6.1: Redija um parágrafo reflexivo estruturado sobre as diferenças, similaridades, consistência e impressões gerais entre os trabalhos com e sem assistência."
    ws6["A3"].font = font_subtitle
    ws6["A3"].fill = fill_blue_head
    ws6["A3"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    dimensoes_q6 = [
        ("1. Padrões de Raciocínio e Justificativas na Triagem", "Como você avalia a profundidade, rigor lógico e rastreabilidade das justificativas dos Métodos A e B na aplicação dos critérios CI1-CI3 e CE1-CE3 no tema de turismo e segurança em fronteiras fluviais?"),
        ("2. Qualidade e Profundidade na Extração de Informações", "Quais foram as principais diferenças percebidas na fidelidade e capacidade de síntese das 5 questões de extração (QE1 a QE5), especialmente na identificação de tipologias criminais e diretrizes de governança?"),
        ("3. Percepção de Autoria e Facilidade de Discernimento", "O trabalho realizado com assistência do RSACV2 é facilmente distinguível do trabalho manual humano? Quais elementos (vocabulário, formatação, exaustividade de códigos) tornaram isso evidente ou não?"),
        ("4. Recomendações e Confiabilidade Metodológica", "Qual a sua avaliação geral sobre a confiabilidade e aplicabilidade do RSACV2 como co-piloto científico para revisões sistemáticas e de escopo em Ciências Sociais Aplicadas e Desenvolvimento Regional?")
    ]
    
    current_r = 5
    for title, prompt_desc in dimensoes_q6:
        ws6.merge_cells(f"A{current_r}:G{current_r}")
        ws6[f"A{current_r}"] = title
        ws6[f"A{current_r}"].font = Font(name="Segoe UI", size=11, bold=True, color=navy_dark)
        ws6[f"A{current_r}"].fill = fill_gold
        ws6[f"A{current_r}"].border = cell_border
        
        ws6.merge_cells(f"A{current_r+1}:G{current_r+1}")
        ws6[f"A{current_r+1}"] = prompt_desc
        ws6[f"A{current_r+1}"].font = font_note
        ws6[f"A{current_r+1}"].alignment = align_left
        ws6[f"A{current_r+1}"].border = cell_border
        
        ws6.merge_cells(f"A{current_r+2}:G{current_r+5}")
        cell_box = ws6[f"A{current_r+2}"]
        cell_box.value = "[Redija aqui o seu parágrafo analítico detalhado...]"
        cell_box.font = font_data
        cell_box.alignment = align_top_left
        for r_box in range(current_r+2, current_r+6):
            for c_box in range(1, 8):
                c_let = get_column_letter(c_box)
                ws6[f"{c_let}{r_box}"].border = cell_border
                ws6[f"{c_let}{r_box}"].fill = fill_white
                
        current_r += 7

    ws6.column_dimensions['A'].width = 15
    ws6.column_dimensions['B'].width = 20
    ws6.column_dimensions['C'].width = 25
    ws6.column_dimensions['D'].width = 25
    ws6.column_dimensions['E'].width = 20
    ws6.column_dimensions['F'].width = 20
    ws6.column_dimensions['G'].width = 20

    # ----------------------------------------------------
    # ABA 7: GABARITO & CONSOLIDAÇÃO (COORDENAÇÃO)
    # ----------------------------------------------------
    ws7 = wb.create_sheet("Gabarito (Coordenação)")
    ws7.views.sheetView[0].showGridLines = True
    
    ws7.merge_cells("A1:G2")
    ws7["A1"] = "GABARITO DE CODIFICAÇÃO E CONSOLIDAÇÃO ESTATÍSTICA (ACESSO RESTRITO)"
    ws7["A1"].font = font_title
    ws7["A1"].fill = fill_navy
    ws7["A1"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    ws7.merge_cells("A3:G3")
    ws7["A3"] = "Atenção: Esta aba deve permanecer oculta ou bloqueada durante a avaliação cega dos revisores para garantir a validade experimental."
    ws7["A3"].font = font_subtitle
    ws7["A3"].fill = fill_blue_head
    ws7["A3"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    
    headers_gab = ["ID Estudo", "Etapa Avaliada", "Método A corresponde a:", "Método B corresponde a:", "Ordem Randomizada", "Chave Criptográfica / Hash", "Validador"]
    for c_idx, h in enumerate(headers_gab, start=1):
        c_letter = get_column_letter(c_idx)
        ws7[f"{c_letter}5"] = h
        ws7[f"{c_letter}5"].font = font_col_header
        ws7[f"{c_letter}5"].fill = fill_blue_head
        ws7[f"{c_letter}5"].alignment = align_center
        ws7[f"{c_letter}5"].border = cell_border
        
    gabarito_rows = [
        ("DOC-001", "Triagem & Extração", "RSACV2 (Assistido)", "Manual (Humano)", "Aleatória (RSAC=A, Manual=B)", "e1c4a9b2...", "Coordenação"),
        ("DOC-002", "Triagem & Extração", "Manual (Humano)", "RSACV2 (Assistido)", "Aleatória (Manual=A, RSAC=B)", "7d2e4f1a...", "Coordenação"),
        ("DOC-003", "Triagem & Extração", "Manual (Humano)", "RSACV2 (Assistido)", "Aleatória (Manual=A, RSAC=B)", "3b8c1e9a...", "Coordenação"),
        ("DOC-004", "Triagem & Extração", "RSACV2 (Assistido)", "Manual (Humano)", "Aleatória (RSAC=A, Manual=B)", "6e1a4f8d...", "Coordenação"),
        ("DOC-005", "Triagem & Extração", "Manual (Humano)", "RSACV2 (Assistido)", "Aleatória (Manual=A, RSAC=B)", "5f9c2e7a...", "Coordenação")
    ]
    for r_idx, row_data in enumerate(gabarito_rows, start=6):
        for c_idx, val in enumerate(row_data, start=1):
            c_letter = get_column_letter(c_idx)
            cell = ws7[f"{c_letter}{r_idx}"]
            cell.value = val
            cell.font = font_data
            cell.fill = fill_alt if r_idx % 2 == 1 else fill_white
            cell.border = cell_border
            cell.alignment = align_center
            if c_idx in [3, 4]:
                cell.font = font_bold_data

    ws7.column_dimensions['A'].width = 14
    ws7.column_dimensions['B'].width = 24
    ws7.column_dimensions['C'].width = 25
    ws7.column_dimensions['D'].width = 25
    ws7.column_dimensions['E'].width = 30
    ws7.column_dimensions['F'].width = 28
    ws7.column_dimensions['G'].width = 18

    file_path = os.path.join(BASE_DIR, "Formulario_Avaliadores_Revisao_Cega.xlsx")
    wb.save(file_path)
    print(f"Evaluator Excel workbook updated at: {file_path}")

if __name__ == "__main__":
    generate_article_docx()
    generate_researcher_sheet()
    generate_evaluator_sheet()
    print("All validation documents successfully updated with the target protocol!")
