"""
Script de Geração da Proposta de Projeto de Pesquisa (Word .docx)
Documento formal para apresentação aos colegas pesquisadores e convite para coautoria.
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

BASE_DIR = r"d:\Downloads\RSACV2\RSACV2\estudo_validacao"
os.makedirs(BASE_DIR, exist_ok=True)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_cell_borders(cell, color="CBD5E1", sz="4", val="single"):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/></w:tcBorders>')
    tcPr.append(tcBorders)

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    for run in h.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(13.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        elif level == 2:
            run.font.size = Pt(11.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2B, 0x54, 0x7E)
        elif level == 3:
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return h

def add_callout(doc, text, title="DESTAQUE METODOLÓGICO", border_color="1B365D", fill_color="F0F4F8"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, fill_color)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r_title = p.add_run(f"[{title}]\n")
    r_title.bold = True
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(9.5)
    r_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    r_text = p.add_run(text)
    r_text.italic = True
    r_text.font.name = "Calibri"
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def generate_project_proposal_docx():
    doc = docx.Document()
    
    # Margens ABNT
    for section in doc.sections:
        section.top_margin = Inches(1.18)    # ~3.0 cm
        section.left_margin = Inches(1.18)   # ~3.0 cm
        section.bottom_margin = Inches(0.79) # ~2.0 cm
        section.right_margin = Inches(0.79)  # ~2.0 cm
        
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Projeto de Pesquisa & Proposta de Coautoria | RSACV2")
        hrun.font.name = "Calibri"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("RSACV2 — Estudo Experimental de Validação • Proposta de Parceria e Coautoria Científica")
        frun.font.name = "Calibri"
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)

    # CABEÇALHO INSTITUCIONAL / CAPA
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_inst.paragraph_format.space_before = Pt(6)
    p_inst.paragraph_format.space_after = Pt(2)
    r_inst1 = p_inst.add_run("PROGRAMA DE PÓS-GRADUAÇÃO EM DESENVOLVIMENTO REGIONAL / CIÊNCIAS SOCIAIS APLICADAS\nGRUPO DE PESQUISA EM METODOLOGIA E TECNOLOGIAS DE REVISÃO SISTEMÁTICA")
    r_inst1.bold = True
    r_inst1.font.size = Pt(10)
    r_inst1.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(14)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("PROJETO DE PESQUISA & CONVITE PARA COAUTORIA CIENTÍFICA")
    r_title.bold = True
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(16)
    r_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(14)
    r_sub = p_sub.add_run("Validação Experimental Pareada Cega da Ferramenta RSACV2 em Revisão de Escopo: Avaliação de Acurácia, Reprodutibilidade e Qualidade Metodológica")
    r_sub.italic = True
    r_sub.font.size = Pt(11.5)
    r_sub.font.color.rgb = RGBColor(0x2B, 0x54, 0x7E)

    # QUADRO DE IDENTIFICAÇÃO DO PROJETO
    tbl_ident = doc.add_table(rows=5, cols=2)
    tbl_ident.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_ident_data = [
        ("Pesquisador Proponente / Coordenador:", "[Seu Nome Completo] — [Titulação / Vínculo Institucional]"),
        ("Linha de Pesquisa:", "Desenvolvimento Regional, Políticas Públicas e Métodos de Síntese de Evidências"),
        ("Protocolo Experimental em Teste:", "'Impactos da Segurança Pública na Operacionalização do Turismo Náutico em Fronteiras Fluviais' (PRISMA-ScR)"),
        ("Previsão de Execução e Escrita:", "[Mês Inicial / 2026] a [Mês Final / 2026] (Duração: ~8 semanas)"),
        ("Periódicos Alvo para Submissão:", "Revistas Qualis A1/A2 em Desenvolvimento Regional, Administração Pública ou Ciência da Informação")
    ]
    col_w_ident = [Inches(2.5), Inches(4.5)]
    for r_idx, (k, v) in enumerate(tbl_ident_data):
        for c_idx, val in enumerate([k, v]):
            cell = tbl_ident.cell(r_idx, c_idx)
            cell.width = col_w_ident[c_idx]
            bg = "F0F4F8" if c_idx == 0 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, 80, 80, 100, 100)
            set_cell_borders(cell, color="CBD5E1", sz="4")
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            if c_idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # CARTA-CONVITE
    tbl_convite = doc.add_table(rows=1, cols=1)
    tbl_convite.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_conv = tbl_convite.cell(0, 0)
    set_cell_background(c_conv, "FEF3C7")
    set_cell_margins(c_conv, top=140, bottom=140, left=180, right=180)
    set_cell_borders(c_conv, color="D97706", sz="8")
    
    p_conv = c_conv.paragraphs[0]
    p_conv.paragraph_format.space_before = Pt(2)
    p_conv.paragraph_format.space_after = Pt(2)
    r_conv_t = p_conv.add_run("CARTA DE CONVITE AOS COLEGAS PESQUISADORES\n")
    r_conv_t.bold = True
    r_conv_t.font.size = Pt(11)
    r_conv_t.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    
    convite_texto = (
        "Prezados(as) Colegas e Pesquisadores(as),\n\n"
        "Gostaria de convidá-los formalmente para integrar a equipe de coautoria e comitê de validação do projeto experimental que visa "
        "testar e validar a ferramenta RSACV2 (Revisão Sistemática Assistida por Computador). "
        "Este estudo possui um desenho metodológico inédito de avaliação cega pareada (duplo-cegamento com trio de revisores independentes), "
        "gerando dados empíricos robustos para publicação de alto impacto na área de Ciências Sociais Aplicadas e Desenvolvimento Regional.\n\n"
        "A participação está estruturada em papéis bem definidos (pesquisador executor ou revisor independente cego, além da coautoria na escrita do artigo), "
        "com previsão de esforço enxuto, instrumental totalmente pronto em planilhas padronizadas e esqueleto do artigo já estruturado. "
        "Abaixo apresento a justificativa, as perguntas de pesquisa, o protocolo em comum e o cronograma de trabalho."
    )
    p_conv.add_run(convite_texto).font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 1. RESUMO EXECUTIVO
    add_styled_heading(doc, "1. RESUMO EXECUTIVO DA PROPOSTA", level=1)
    doc.add_paragraph(
        "A condução de revisões sistemáticas e revisões de escopo (Scoping Reviews) é hoje uma das metodologias mais valorizadas "
        "na pesquisa acadêmica contemporânea. No entanto, as etapas de triagem de centenas de resumos e extração de dados demandam "
        "centenas de horas de trabalho humano, frequentemente sujeitas a fadiga e vieses de interpretação."
    )
    doc.add_paragraph(
        "A ferramenta RSACV2 foi desenvolvida para automatizar e assistir essas etapas com ancoragem estrita, transparência em justificativas "
        "e respeito às limitações de motores de busca nacionais e internacionais (BDTD e SciELO). "
        "Para que seu uso seja cientificamente legitimado perante a comunidade acadêmica e bancas avaliadoras, propomos a realização de um "
        "estudo experimental controlado que compare, de forma rigorosa e cega, os resultados gerados pela ferramenta versus o método manual humano."
    )

    # 2. JUSTIFICATIVA E RELEVÂNCIA
    add_styled_heading(doc, "2. JUSTIFICATIVA E RELEVÂNCIA DO ESTUDO", level=1)
    doc.add_paragraph(
        "A literatura recente sobre automação de síntese de evidências (BARKER et al., 2024; MARSHALL; WALLACE, 2019) aponta que a validação "
        "de novas ferramentas exige experimentos de concordância inter-avaliadores (Inter-rater reliability) com índices formais (Kappa de Cohen e Fleiss). "
        "Entretanto, a quase totalidade dos benchmarks existentes concentra-se em áreas biomédicas/clínicas."
    )
    doc.add_paragraph(
        "Este projeto diferencia-se por aplicar a validação no domínio de Ciências Sociais Aplicadas e Desenvolvimento Regional, "
        "utilizando como objeto de prova o protocolo: 'Impactos da Segurança Pública na Operacionalização do Turismo Náutico em Fronteiras Fluviais'."
    )

    # 3. QUESTÕES DE PESQUISA
    add_styled_heading(doc, "3. OBJETIVOS E QUESTÕES DE PESQUISA (RESEARCH QUESTIONS)", level=1)
    doc.add_paragraph("O projeto responderá a seis perguntas de pesquisa fundamentais:")

    rq_list = [
        ("RQ1 (Coleta e Recuperação)", "Usando o mesmo protocolo de busca, em uma mesma base, os trabalhos encontrados pelo RSACV2 e pelo método manual são os mesmos em presença e ordem de aparição? (Critério 1.1: SIM, NÃO POR PRESENÇA, NÃO POR ORDEM, NÃO INTEGRAL)"),
        ("RQ2 (Concordância na Triagem)", "Em uma mesma lista de critérios de inclusão (CI1-CI3) e exclusão (CE1-CE3), as decisões de triagem serão as mesmas? E a qualidade da fundamentação das justificativas?"),
        ("RQ3 (Preferência em Divergências)", "Em caso de divergências de marcações e justificativas, qual alternativa é preferida por um trio de revisores independentes em uma análise cega?"),
        ("RQ4 (Discernibilidade / Teste de Turing)", "Os revisores independentes serão capazes de diferenciar qual resposta foi gerada com auxílio da ferramenta? Sob qual grau de certeza (escala 1 a 5) e quais justificativas?"),
        ("RQ5 (Qualidade na Extração de Dados)", "Como se comparam as extrações de dados (QE1 a QE5) nos dois métodos sob avaliação cega em termos de profundidade, completude e concisão?"),
        ("RQ6 (Síntese Qualitativa Global)", "Quais observações e padrões de similaridade/diferença os revisores declaram sobre a confiabilidade e robustez do trabalho assistido?")
    ]
    for rq_id, rq_desc in rq_list:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"• {rq_id}: ")
        r.bold = True
        r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        p.add_run(rq_desc)

    # 4. O PROTOCOLO OBJETO DE VALIDAÇÃO
    add_styled_heading(doc, "4. O PROTOCOLO OBJETO DE TESTE (TURISMO NÁUTICO EM FRONTEIRAS FLUVIAIS)", level=1)
    doc.add_paragraph(
        "O protocolo adotado em comum para todos os participantes está registrado no sistema RSACV2 (PRISMA-ScR) e aborda as problemáticas de "
        "segurança pública, criminalidade transfronteiriça e fiscalização em hidrovias e bacias de fronteira (Amazônica, Bacia do Prata, etc.) "
        "e seus impactos na atividade turística náutica."
    )

    # Tabela com Síntese do Protocolo
    tbl_proto_sum = doc.add_table(rows=5, cols=2)
    tbl_proto_sum.alignment = WD_TABLE_ALIGNMENT.CENTER
    proto_items = [
        ("Bases e Filtros:", "BDTD (Teses e Dissertações) e SciELO (Artigos de Periódicos); Idiomas: PT, EN, ES."),
        ("Descritores de Busca (Pares):", "5 pares em PT, 5 em EN e 5 em ES (ex: 'turismo náutico' AND 'fronteira'; 'segurança pública' AND 'fronteira fluvial'; 'turismo' AND 'tríplice fronteira')."),
        ("Critérios de Inclusão (CI):", "CI1 (Turismo/Navegação Fluvial em Fronteira); CI2 (Segurança Pública/Governança de Bacia); CI3 (Artigo, Tese ou Dissertação completa)."),
        ("Critérios de Exclusão (CE):", "CE1 (Transporte Marítimo Oceânico de Alto-Mar); CE2 (Segurança Urbana/Rural Desvinculada de Hidrovias); CE3 (Editorial/Resumo sem método)."),
        ("Questões de Extração (QE):", "QE1 (Localização & Bacia); QE2 (Tipologia de Crimes/Ilícitos); QE3 (Impactos no Turismo Náutico); QE4 (Governança e Políticas Recomendadas); QE5 (Metodologia).")
    ]
    for r_idx, (k, v) in enumerate(proto_items):
        for c_idx, val in enumerate([k, v]):
            cell = tbl_proto_sum.cell(r_idx, c_idx)
            cell.width = col_w_ident[c_idx]
            bg = "F0F4F8" if c_idx == 0 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, 60, 60, 80, 80)
            set_cell_borders(cell, color="CBD5E1", sz="4")
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            if c_idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 5. ESTRUTURA DA EQUIPE E PAPÉIS
    add_styled_heading(doc, "5. ESTRUTURA DA EQUIPE, PAPÉIS E ATRIBUIÇÕES", level=1)
    doc.add_paragraph(
        "Para garantir o cegamento metodológico e a viabilidade operacional, o trabalho será dividido em 3 núcleos cooperativos:"
    )

    papeis = [
        ("Núcleo 1: Coordenação Geral e Cegamento (1 a 2 pesquisadores)",
         "Responsável pela gestão do cronograma, exportação dos dados brutos do RSACV2, anonimização e aleatorização dos formulários (Método A vs. Método B), guarda do gabarito criptografado e compilação dos testes estatísticos de concordância."),
        ("Núcleo 2: Pesquisadores Executores do Braço Manual (1 a 2 pesquisadores)",
         "Responsáveis por realizar a busca independente nas bases, efetuar a triagem na planilha padronizada e preencher as 5 questões de extração para a amostra de artigos incluídos, sem conhecimento das respostas geradas pelo RSACV2."),
        ("Núcleo 3: Comitê de Avaliadores Independentes / Trio Cego (3 pesquisadores)",
         "Responsáveis por receber a planilha cega codificada, comparar os Métodos A e B, julgar a coleta (1.1), indicar preferências nas divergências (3.1), realizar o teste de percepção/Turing (4.1), avaliar a extração (5.1/5.2) e redigir o parecer qualitativo (6.1)."),
        ("Todos os Integrantes: Redação Científica e Coautoria",
         "Participação na contextualização teórica, discussão dos resultados, revisão crítica e aprovação final do manuscrito para submissão.")
    ]
    for p_title, p_desc in papeis:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"• {p_title}:\n  ")
        r.bold = True
        r.font.color.rgb = RGBColor(0x2B, 0x54, 0x7E)
        p.add_run(p_desc)

    # 6. CRONOGRAMA
    add_styled_heading(doc, "6. CRONOGRAMA DE EXECUÇÃO (8 SEMANAS)", level=1)
    
    tbl_cron = doc.add_table(rows=6, cols=3)
    tbl_cron.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_cron = ["Semana(s)", "Atividade / Etapa", "Entregável / Meta"]
    col_w_cron = [Inches(1.5), Inches(3.2), Inches(2.3)]
    
    for c_idx, h in enumerate(headers_cron):
        cell = tbl_cron.cell(0, c_idx)
        cell.width = col_w_cron[c_idx]
        set_cell_background(cell, "1B365D")
        set_cell_margins(cell, 80, 80, 100, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)

    cron_data = [
        ("Semanas 1-2", "Alinhamento da equipe, distribuição dos formulários e execução da coleta/triagem manual e assistida", "Planilhas dos pesquisadores preenchidas"),
        ("Semana 3", "Anonimização, codificação (Método A vs. B) e envio ao Trio de Avaliadores", "Planilha cega distribuída aos revisores"),
        ("Semanas 4-5", "Avaliação cega independente pelo Trio de Revisores (RQ1 a RQ6)", "Formulários dos revisores preenchidos"),
        ("Semana 6", "Abertura do gabarito, compilação estatística (Kappa de Cohen/Fleiss, Teste de Turing, Likert)", "Tabelas de resultados consolidadas"),
        ("Semanas 7-8", "Redação conjunta das seções de Discussão/Conclusão, revisão final e submissão do artigo", "Manuscrito finalizado e submetido")
    ]
    for r_idx, row in enumerate(cron_data, start=1):
        for c_idx, val in enumerate(row):
            cell = tbl_cron.cell(r_idx, c_idx)
            cell.width = col_w_cron[c_idx]
            set_cell_background(cell, "F0F4F8" if r_idx % 2 == 1 else "FFFFFF")
            set_cell_margins(cell, 60, 60, 80, 80)
            set_cell_borders(cell, color="CBD5E1", sz="4")
            p = cell.paragraphs[0]
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(val).font.size = Pt(8.5)
            else:
                p.add_run(val).font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 7. CRITÉRIOS DE COAUTORIA
    add_styled_heading(doc, "7. CRITÉRIOS DE COAUTORIA E INTEGRIDADE CIENTÍFICA", level=1)
    doc.add_paragraph(
        "A coautoria será regida estritamente pelas diretrizes éticas internacionais do Committee on Publication Ethics (COPE) "
        "e do International Committee of Medical Journal Editors (ICMJE). Todos os participantes que desempenharem as atividades de execução, "
        "avaliação ou redação terão sua coautoria formalmente assegurada no manuscrito final."
    )

    # 8. TERMO DE ADESÃO
    add_styled_heading(doc, "8. FICHA DE ADESÃO / INDICAÇÃO DE INTERESSE", level=1)
    doc.add_paragraph("Favor preencher os dados abaixo para formalização da equipe de pesquisa:")

    tbl_adesao = doc.add_table(rows=6, cols=2)
    tbl_adesao.alignment = WD_TABLE_ALIGNMENT.CENTER
    adesao_fields = [
        ("Nome Completo:", ""),
        ("E-mail e Telefone / WhatsApp:", ""),
        ("Instituição / Programa de Pós-Graduação:", ""),
        ("Titulação / Área de Atuação:", ""),
        ("Papel de Preferência:", "( ) Pesquisador Executor Manual   ( ) Revisor Cego Independente   ( ) Redator/Analista"),
        ("Assinatura / Concordância:", "Data: ____/____/2026   Assinatura: ___________________________")
    ]
    for r_idx, (k, v) in enumerate(adesao_fields):
        for c_idx, val in enumerate([k, v]):
            cell = tbl_adesao.cell(r_idx, c_idx)
            cell.width = Inches(2.5) if c_idx == 0 else Inches(4.5)
            set_cell_background(cell, "F0F4F8" if c_idx == 0 else "FFFFFF")
            set_cell_margins(cell, 70, 70, 90, 90)
            set_cell_borders(cell, color="CBD5E1", sz="4")
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            if c_idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    doc_path = os.path.join(BASE_DIR, "Projeto_de_Pesquisa_Proposta_Validacao_RSACV2.docx")
    doc.save(doc_path)
    print(f"Project Proposal Word document created at: {doc_path}")

if __name__ == "__main__":
    generate_project_proposal_docx()
